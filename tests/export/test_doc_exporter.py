"""Tests for src/export/doc_exporter.py - Individual document export functionality."""

import base64
import tempfile
import zipfile
from datetime import datetime
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest

from src.export.doc_exporter import (
    slugify,
    BaseDocExporter,
    DocPDFExporter,
    DocWordExporter,
    DocHTMLExporter,
    DocMarkdownExporter,
    create_doc_exporter,
)


class TestSlugify:
    """Tests for the slugify utility function."""

    def test_slugify_lowercase(self):
        """Test that slugify converts to lowercase."""
        assert slugify("HELLO") == "hello"
        assert slugify("Hello World") == "hello-world"

    def test_slugify_replaces_spaces(self):
        """Test that slugify replaces spaces with hyphens."""
        assert slugify("hello world") == "hello-world"
        assert slugify("hello  world") == "hello-world"

    def test_slugify_replaces_underscores(self):
        """Test that slugify replaces underscores with hyphens."""
        assert slugify("hello_world") == "hello-world"

    def test_slugify_removes_special_chars(self):
        """Test that slugify removes special characters."""
        assert slugify("hello!@#$%world") == "helloworld"
        assert slugify("test (example)") == "test-example"

    def test_slugify_collapses_hyphens(self):
        """Test that slugify collapses multiple hyphens."""
        assert slugify("hello---world") == "hello-world"
        assert slugify("a - b - c") == "a-b-c"

    def test_slugify_strips_hyphens(self):
        """Test that slugify strips leading/trailing hyphens."""
        assert slugify("-hello-") == "hello"
        assert slugify("---test---") == "test"

    def test_slugify_with_numbers(self):
        """Test that slugify preserves numbers."""
        assert slugify("chapter1") == "chapter1"
        assert slugify("Chapter 2 - Advanced") == "chapter-2-advanced"


class TestBaseDocExporter:
    """Tests for BaseDocExporter abstract class."""

    def test_base_exporter_is_abstract(self):
        """Test that BaseDocExporter cannot be instantiated directly."""
        with pytest.raises(TypeError):
            BaseDocExporter("user-id", "doc-id")

    def test_file_extension_default(self):
        """Test that file_extension has empty default."""
        assert BaseDocExporter.file_extension == ""


class TestDocPDFExporter:
    """Tests for DocPDFExporter class."""

    def test_file_extension(self):
        """Test that file extension is 'pdf'."""
        assert DocPDFExporter.file_extension == "pdf"

    def test_default_css_exists(self):
        """Test that DEFAULT_CSS is defined."""
        assert hasattr(DocPDFExporter, "DEFAULT_CSS")
        assert len(DocPDFExporter.DEFAULT_CSS) > 0

    def test_default_css_has_page_settings(self):
        """Test that DEFAULT_CSS includes @page rule."""
        assert "@page" in DocPDFExporter.DEFAULT_CSS
        assert "A4" in DocPDFExporter.DEFAULT_CSS

    @pytest.fixture
    def pdf_exporter(self, tmp_path):
        """Create a DocPDFExporter for testing."""
        exporter = DocPDFExporter.__new__(DocPDFExporter)
        storage = MagicMock()
        storage.docs_dir = tmp_path / "docs"
        storage.docs_dir.mkdir(parents=True)
        storage.get_doc_content = MagicMock(return_value="# Test Manual\n\nContent here.")
        storage.get_doc_metadata = MagicMock(return_value={"title": "Test Manual"})
        exporter.user_storage = storage
        exporter.user_id = "test-user"
        exporter.doc_id = "test-doc"
        return exporter

    def test_get_output_path_generates_path(self, pdf_exporter, tmp_path):
        """Test that _get_output_path generates a valid path."""
        path = pdf_exporter._get_output_path(None, "en")

        assert path.endswith(".pdf")
        assert "test-doc" in path
        assert "en" in path

    def test_get_output_path_uses_provided_path(self, pdf_exporter, tmp_path):
        """Test that _get_output_path uses provided path."""
        provided = str(tmp_path / "custom.pdf")
        path = pdf_exporter._get_output_path(provided, "en")

        assert path == provided


class TestDocWordExporter:
    """Tests for DocWordExporter class."""

    def test_file_extension(self):
        """Test that file extension is 'docx'."""
        assert DocWordExporter.file_extension == "docx"

    @pytest.fixture
    def word_exporter(self, tmp_path):
        """Create a DocWordExporter for testing."""
        exporter = DocWordExporter.__new__(DocWordExporter)
        storage = MagicMock()
        storage.docs_dir = tmp_path / "docs"
        storage.docs_dir.mkdir(parents=True)
        storage.get_doc_content = MagicMock(return_value="# Test Manual\n\nContent here.")
        storage.get_doc_metadata = MagicMock(return_value={"title": "Test Manual"})
        exporter.user_storage = storage
        exporter.user_id = "test-user"
        exporter.doc_id = "test-doc"
        return exporter

    def test_export_creates_docx(self, word_exporter, tmp_path):
        """Test that export creates a .docx file."""
        output_path = str(tmp_path / "output.docx")

        with patch("src.export.doc_exporter.strip_semantic_tags", return_value="# Test"):
            result = word_exporter.export(output_path=output_path)

        assert Path(result).exists()
        assert result.endswith(".docx")

    def test_add_markdown_to_doc_headers(self, word_exporter):
        """Test parsing markdown headers."""
        from docx import Document
        doc = Document()

        content = "# H1\n## H2\n### H3\n#### H4"
        word_exporter._add_markdown_to_doc(doc, content)

        # Check that headings were added
        headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert len(headings) >= 4

    def test_add_markdown_to_doc_lists(self, word_exporter):
        """Test parsing markdown lists."""
        from docx import Document
        doc = Document()

        content = "- Item 1\n- Item 2\n1. Num 1\n2. Num 2"
        word_exporter._add_markdown_to_doc(doc, content)

        # Check that list items were added
        assert len(doc.paragraphs) >= 4


class TestDocHTMLExporter:
    """Tests for DocHTMLExporter class."""

    def test_file_extension(self):
        """Test that file extension is 'html'."""
        assert DocHTMLExporter.file_extension == "html"

    def test_default_css_exists(self):
        """Test that DEFAULT_CSS is defined."""
        assert hasattr(DocHTMLExporter, "DEFAULT_CSS")
        assert len(DocHTMLExporter.DEFAULT_CSS) > 0

    def test_default_css_has_dark_mode(self):
        """Test that DEFAULT_CSS includes dark mode."""
        assert "prefers-color-scheme: dark" in DocHTMLExporter.DEFAULT_CSS

    @pytest.fixture
    def html_exporter(self, tmp_path):
        """Create a DocHTMLExporter for testing."""
        exporter = DocHTMLExporter.__new__(DocHTMLExporter)
        storage = MagicMock()
        storage.docs_dir = tmp_path / "docs"
        storage.docs_dir.mkdir(parents=True)
        storage.get_doc_content = MagicMock(return_value="# Test Manual\n\nContent here.")
        storage.get_doc_metadata = MagicMock(return_value={"title": "Test Manual"})
        exporter.user_storage = storage
        exporter.user_id = "test-user"
        exporter.doc_id = "test-doc"
        return exporter

    def test_export_creates_html(self, html_exporter, tmp_path):
        """Test that export creates an .html file."""
        output_path = str(tmp_path / "output.html")

        with patch("src.export.doc_exporter.strip_semantic_tags", return_value="# Test"):
            result = html_exporter.export(output_path=output_path)

        assert Path(result).exists()
        assert result.endswith(".html")

        with open(result, "r") as f:
            content = f.read()
        assert "<!DOCTYPE html>" in content

    def test_export_embed_images_default(self, html_exporter, tmp_path):
        """Test that images are embedded by default."""
        output_path = str(tmp_path / "embed.html")

        with patch("src.export.doc_exporter.strip_semantic_tags", return_value="# Test"):
            html_exporter._embed_images_as_base64 = MagicMock(return_value="# Test")
            result = html_exporter.export(output_path=output_path, embed_images=True)

        html_exporter._embed_images_as_base64.assert_called()

    def test_export_no_embed_images(self, html_exporter, tmp_path):
        """Test that image embedding can be disabled."""
        output_path = str(tmp_path / "no_embed.html")

        with patch("src.export.doc_exporter.strip_semantic_tags", return_value="# Test"):
            html_exporter._embed_images_as_base64 = MagicMock()
            html_exporter._fix_image_paths = MagicMock(return_value="# Test")
            result = html_exporter.export(output_path=output_path, embed_images=False)

        html_exporter._embed_images_as_base64.assert_not_called()
        html_exporter._fix_image_paths.assert_called()


class TestDocMarkdownExporter:
    """Tests for DocMarkdownExporter class."""

    def test_file_extension(self):
        """Test that file extension is 'zip'."""
        assert DocMarkdownExporter.file_extension == "zip"

    @pytest.fixture
    def md_exporter(self, tmp_path):
        """Create a DocMarkdownExporter for testing."""
        exporter = DocMarkdownExporter.__new__(DocMarkdownExporter)
        storage = MagicMock()
        storage.docs_dir = tmp_path / "docs"
        storage.docs_dir.mkdir(parents=True)
        storage.get_doc_content = MagicMock(return_value="# Test\n\n![Image](screenshots/test.png)")
        exporter.user_storage = storage
        exporter.user_id = "test-user"
        exporter.doc_id = "test-doc"
        return exporter

    def test_export_creates_zip(self, md_exporter, tmp_path):
        """Test that export creates a .zip file."""
        output_path = str(tmp_path / "output.zip")

        # Create screenshots directory and image
        screenshots_dir = md_exporter.user_storage.docs_dir / "test-doc" / "screenshots"
        screenshots_dir.mkdir(parents=True)
        from PIL import Image
        img = Image.new("RGB", (10, 10), "red")
        img.save(screenshots_dir / "test.png")

        with patch("src.export.doc_exporter.strip_semantic_tags", return_value="# Test\n\n![Image](screenshots/test.png)"):
            result = md_exporter.export(output_path=output_path)

        assert Path(result).exists()
        assert result.endswith(".zip")

    def test_zip_contains_markdown(self, md_exporter, tmp_path):
        """Test that ZIP contains markdown file."""
        output_path = str(tmp_path / "output.zip")

        screenshots_dir = md_exporter.user_storage.docs_dir / "test-doc" / "screenshots"
        screenshots_dir.mkdir(parents=True)

        with patch("src.export.doc_exporter.strip_semantic_tags", return_value="# Test"):
            result = md_exporter.export(output_path=output_path)

        with zipfile.ZipFile(result, "r") as zf:
            names = zf.namelist()
            assert any(name.endswith(".md") for name in names)

    def test_zip_contains_images(self, md_exporter, tmp_path):
        """Test that ZIP contains referenced images."""
        output_path = str(tmp_path / "output.zip")

        screenshots_dir = md_exporter.user_storage.docs_dir / "test-doc" / "screenshots"
        screenshots_dir.mkdir(parents=True)
        from PIL import Image
        img = Image.new("RGB", (10, 10), "red")
        img.save(screenshots_dir / "test.png")

        with patch("src.export.doc_exporter.strip_semantic_tags", return_value="# Test\n\n![Image](screenshots/test.png)"):
            result = md_exporter.export(output_path=output_path)

        with zipfile.ZipFile(result, "r") as zf:
            names = zf.namelist()
            assert any("images/test.png" in name for name in names)


class TestCreateDocExporter:
    """Tests for create_doc_exporter factory function."""

    def test_create_pdf_exporter(self, tmp_path):
        """Test creating PDF exporter."""
        with patch("src.export.doc_exporter.UserStorage") as mock_storage:
            mock_storage.return_value.docs_dir = tmp_path
            (tmp_path / "test-doc").mkdir()

            exporter = create_doc_exporter("user-id", "test-doc", "pdf")
            assert isinstance(exporter, DocPDFExporter)

    def test_create_word_exporter(self, tmp_path):
        """Test creating Word exporter."""
        with patch("src.export.doc_exporter.UserStorage") as mock_storage:
            mock_storage.return_value.docs_dir = tmp_path
            (tmp_path / "test-doc").mkdir()

            exporter = create_doc_exporter("user-id", "test-doc", "word")
            assert isinstance(exporter, DocWordExporter)

    def test_create_docx_exporter(self, tmp_path):
        """Test creating Word exporter with 'docx' format."""
        with patch("src.export.doc_exporter.UserStorage") as mock_storage:
            mock_storage.return_value.docs_dir = tmp_path
            (tmp_path / "test-doc").mkdir()

            exporter = create_doc_exporter("user-id", "test-doc", "docx")
            assert isinstance(exporter, DocWordExporter)

    def test_create_html_exporter(self, tmp_path):
        """Test creating HTML exporter."""
        with patch("src.export.doc_exporter.UserStorage") as mock_storage:
            mock_storage.return_value.docs_dir = tmp_path
            (tmp_path / "test-doc").mkdir()

            exporter = create_doc_exporter("user-id", "test-doc", "html")
            assert isinstance(exporter, DocHTMLExporter)

    def test_create_markdown_exporter(self, tmp_path):
        """Test creating Markdown exporter."""
        with patch("src.export.doc_exporter.UserStorage") as mock_storage:
            mock_storage.return_value.docs_dir = tmp_path
            (tmp_path / "test-doc").mkdir()

            exporter = create_doc_exporter("user-id", "test-doc", "markdown")
            assert isinstance(exporter, DocMarkdownExporter)

    def test_create_md_exporter(self, tmp_path):
        """Test creating Markdown exporter with 'md' format."""
        with patch("src.export.doc_exporter.UserStorage") as mock_storage:
            mock_storage.return_value.docs_dir = tmp_path
            (tmp_path / "test-doc").mkdir()

            exporter = create_doc_exporter("user-id", "test-doc", "md")
            assert isinstance(exporter, DocMarkdownExporter)

    def test_create_exporter_case_insensitive(self, tmp_path):
        """Test that format is case insensitive."""
        with patch("src.export.doc_exporter.UserStorage") as mock_storage:
            mock_storage.return_value.docs_dir = tmp_path
            (tmp_path / "test-doc").mkdir()

            exporter1 = create_doc_exporter("user-id", "test-doc", "PDF")
            exporter2 = create_doc_exporter("user-id", "test-doc", "Pdf")

            assert isinstance(exporter1, DocPDFExporter)
            assert isinstance(exporter2, DocPDFExporter)

    def test_create_exporter_invalid_format(self, tmp_path):
        """Test that invalid format raises ValueError."""
        with patch("src.export.doc_exporter.UserStorage") as mock_storage:
            mock_storage.return_value.docs_dir = tmp_path
            (tmp_path / "test-doc").mkdir()

            with pytest.raises(ValueError) as exc_info:
                create_doc_exporter("user-id", "test-doc", "invalid")

            assert "Unsupported export format" in str(exc_info.value)


class TestExportSingleDoc:
    """Tests for exporting single documents."""

    @pytest.fixture
    def html_exporter(self, tmp_path):
        """Create DocHTMLExporter for single doc testing."""
        exporter = DocHTMLExporter.__new__(DocHTMLExporter)
        storage = MagicMock()
        storage.docs_dir = tmp_path / "docs"
        storage.docs_dir.mkdir(parents=True)
        storage.get_doc_content = MagicMock(return_value="# Test Manual\n\nContent.")
        storage.get_doc_metadata = MagicMock(return_value={"title": "Test"})
        exporter.user_storage = storage
        exporter.user_id = "test-user"
        exporter.doc_id = "test-doc"
        return exporter

    def test_export_single_doc_basic(self, html_exporter, tmp_path):
        """Test basic single document export."""
        output_path = str(tmp_path / "single.html")

        with patch("src.export.doc_exporter.strip_semantic_tags", return_value="# Test"):
            result = html_exporter.export(output_path=output_path)

        assert Path(result).exists()

    def test_export_single_doc_with_language(self, html_exporter, tmp_path):
        """Test single document export with language parameter."""
        output_path = str(tmp_path / "spanish.html")

        with patch("src.export.doc_exporter.strip_semantic_tags", return_value="# Test"):
            result = html_exporter.export(output_path=output_path, language="es")

        with open(result, "r") as f:
            content = f.read()

        # Check language is in the HTML
        assert 'lang="es"' in content


class TestExportWithOptions:
    """Tests for export with various options."""

    @pytest.fixture
    def html_exporter(self, tmp_path):
        """Create DocHTMLExporter for options testing."""
        exporter = DocHTMLExporter.__new__(DocHTMLExporter)
        storage = MagicMock()
        storage.docs_dir = tmp_path / "docs"
        storage.docs_dir.mkdir(parents=True)
        storage.get_doc_content = MagicMock(return_value="# Test\n\n![Image](screenshots/test.png)")
        storage.get_doc_metadata = MagicMock(return_value={"title": "Test"})
        exporter.user_storage = storage
        exporter.user_id = "test-user"
        exporter.doc_id = "test-doc"
        return exporter

    def test_export_with_embed_images_true(self, html_exporter, tmp_path):
        """Test export with embed_images=True."""
        output_path = str(tmp_path / "embedded.html")

        # Create test image
        screenshots_dir = html_exporter.user_storage.docs_dir / "test-doc" / "screenshots"
        screenshots_dir.mkdir(parents=True)
        from PIL import Image
        img = Image.new("RGB", (10, 10), "red")
        img.save(screenshots_dir / "test.png")

        with patch("src.export.doc_exporter.strip_semantic_tags", return_value="# Test\n\n![Image](screenshots/test.png)"):
            result = html_exporter.export(output_path=output_path, embed_images=True)

        with open(result, "r") as f:
            content = f.read()

        # Should contain base64 data
        assert "data:" in content or "base64" in content

    def test_export_with_embed_images_false(self, html_exporter, tmp_path):
        """Test export with embed_images=False."""
        output_path = str(tmp_path / "not_embedded.html")

        with patch("src.export.doc_exporter.strip_semantic_tags", return_value="# Test"):
            result = html_exporter.export(output_path=output_path, embed_images=False)

        assert Path(result).exists()


class TestImagePathFixing:
    """Tests for image path handling in exports."""

    @pytest.fixture
    def exporter(self, tmp_path):
        """Create exporter for image path tests."""
        exporter = DocHTMLExporter.__new__(DocHTMLExporter)
        storage = MagicMock()
        storage.docs_dir = tmp_path / "docs"
        storage.docs_dir.mkdir(parents=True)
        exporter.user_storage = storage
        exporter.doc_id = "test-doc"
        return exporter

    def test_fix_image_paths_relative(self, exporter, tmp_path):
        """Test fixing relative image paths."""
        # Create screenshot directory and file
        screenshots_dir = exporter.user_storage.docs_dir / "test-doc" / "screenshots"
        screenshots_dir.mkdir(parents=True)
        (screenshots_dir / "test.png").touch()

        content = "![Alt](../screenshots/test.png)"
        result = exporter._fix_image_paths(content)

        assert "file://" in result

    def test_fix_image_paths_screenshots_prefix(self, exporter, tmp_path):
        """Test fixing screenshots/ prefixed paths."""
        screenshots_dir = exporter.user_storage.docs_dir / "test-doc" / "screenshots"
        screenshots_dir.mkdir(parents=True)
        (screenshots_dir / "test.png").touch()

        content = "![Alt](screenshots/test.png)"
        result = exporter._fix_image_paths(content)

        assert "file://" in result

    def test_fix_image_paths_preserves_missing(self, exporter):
        """Test that missing images paths are preserved."""
        content = "![Alt](nonexistent.png)"
        result = exporter._fix_image_paths(content)

        assert "nonexistent.png" in result


class TestEmbedImagesAsBase64:
    """Tests for base64 image embedding."""

    @pytest.fixture
    def html_exporter(self, tmp_path):
        """Create DocHTMLExporter for base64 tests."""
        exporter = DocHTMLExporter.__new__(DocHTMLExporter)
        storage = MagicMock()
        storage.docs_dir = tmp_path / "docs"
        storage.docs_dir.mkdir(parents=True)
        exporter.user_storage = storage
        exporter.doc_id = "test-doc"
        return exporter

    def test_embed_png_image(self, html_exporter, tmp_path):
        """Test embedding PNG image."""
        screenshots_dir = html_exporter.user_storage.docs_dir / "test-doc" / "screenshots"
        screenshots_dir.mkdir(parents=True)
        from PIL import Image
        img = Image.new("RGB", (10, 10), "red")
        img.save(screenshots_dir / "test.png")

        content = "![Alt](screenshots/test.png)"
        result = html_exporter._embed_images_as_base64(content)

        assert "data:image/png;base64," in result

    def test_embed_jpg_image(self, html_exporter, tmp_path):
        """Test embedding JPG image."""
        screenshots_dir = html_exporter.user_storage.docs_dir / "test-doc" / "screenshots"
        screenshots_dir.mkdir(parents=True)
        from PIL import Image
        img = Image.new("RGB", (10, 10), "blue")
        img.save(screenshots_dir / "test.jpg", "JPEG")

        content = "![Alt](screenshots/test.jpg)"
        result = html_exporter._embed_images_as_base64(content)

        assert "data:image/jpeg;base64," in result

    def test_embed_gif_image(self, html_exporter, tmp_path):
        """Test embedding GIF image."""
        screenshots_dir = html_exporter.user_storage.docs_dir / "test-doc" / "screenshots"
        screenshots_dir.mkdir(parents=True)
        from PIL import Image
        img = Image.new("RGB", (10, 10), "green")
        img.save(screenshots_dir / "test.gif", "GIF")

        content = "![Alt](screenshots/test.gif)"
        result = html_exporter._embed_images_as_base64(content)

        assert "data:image/gif;base64," in result

    def test_embed_skips_external_urls(self, html_exporter):
        """Test that external URLs are not embedded."""
        content = "![Alt](https://example.com/image.png)"
        result = html_exporter._embed_images_as_base64(content)

        assert "https://example.com/image.png" in result
        assert "data:image" not in result

    def test_embed_skips_data_urls(self, html_exporter):
        """Test that data URLs are not re-embedded."""
        content = "![Alt](data:image/png;base64,ABC123)"
        result = html_exporter._embed_images_as_base64(content)

        assert "data:image/png;base64,ABC123" in result

    def test_embed_preserves_missing_images(self, html_exporter):
        """Test that missing images are preserved as-is."""
        content = "![Alt](screenshots/nonexistent.png)"
        result = html_exporter._embed_images_as_base64(content)

        assert "screenshots/nonexistent.png" in result
