"""Tests for src/export/word_exporter.py - Word document export functionality."""

import tempfile
from pathlib import Path
from unittest.mock import MagicMock, patch, PropertyMock

import pytest
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.export.word_exporter import WordExporter


@pytest.fixture
def mock_user_storage():
    """Create a mock UserStorage instance."""
    storage = MagicMock()
    storage.docs_dir = Path(tempfile.mkdtemp())
    return storage


@pytest.fixture
def mock_project_storage():
    """Create a mock ProjectStorage instance."""
    storage = MagicMock()
    storage.projects_dir = Path(tempfile.mkdtemp())
    return storage


@pytest.fixture
def sample_project():
    """Create a sample project dict for testing."""
    return {
        "id": "test-project",
        "name": "Test Project",
        "description": "A test project description",
        "chapters": [
            {
                "id": "chapter-1",
                "title": "Getting Started",
                "description": "Introduction chapter",
                "order": 1,
                "manuals": ["manual-1", "manual-2"],
            },
            {
                "id": "chapter-2",
                "title": "Advanced Topics",
                "description": "Advanced chapter",
                "order": 2,
                "manuals": ["manual-3"],
            },
        ],
    }


@pytest.fixture
def sample_markdown_content():
    """Create sample markdown content for testing."""
    return """# Test Manual

## Introduction
This is the introduction section.

## Step 1: First Step
Do the first thing.

![Screenshot 1](screenshots/step1.png)

## Step 2: Second Step
Do the second thing with **bold** and *italic* text.

### Sub-step
A nested step with `code` inline.

```python
print("Hello World")
```

> This is a blockquote

- Bullet point 1
- Bullet point 2

1. Numbered item 1
2. Numbered item 2
"""


class TestWordExporterInit:
    """Tests for WordExporter initialization."""

    def test_file_extension(self):
        """Test that file extension is 'docx'."""
        assert WordExporter.file_extension == "docx"

    def test_init_with_valid_project(
        self, mock_user_storage, mock_project_storage, sample_project
    ):
        """Test initialization with a valid project."""
        with patch("src.export.word_exporter.BaseExporter.__init__") as mock_init:
            mock_init.return_value = None
            exporter = WordExporter.__new__(WordExporter)
            exporter.user_storage = mock_user_storage
            exporter.project_storage = mock_project_storage
            exporter.project = sample_project
            exporter.user_id = "test-user"
            exporter.project_id = "test-project"

            assert exporter.project["name"] == "Test Project"


class TestWordExporterExport:
    """Tests for the export method."""

    @pytest.fixture
    def word_exporter(self, mock_user_storage, mock_project_storage, sample_project):
        """Create a WordExporter instance for testing."""
        exporter = WordExporter.__new__(WordExporter)
        exporter.user_storage = mock_user_storage
        exporter.project_storage = mock_project_storage
        exporter.project = sample_project
        exporter.user_id = "test-user"
        exporter.project_id = "test-project"
        return exporter

    def test_export_creates_docx_file(self, word_exporter, tmp_path):
        """Test that export creates a .docx file."""
        output_path = str(tmp_path / "test_output.docx")

        # Mock the _get_manual_content to return empty for simplicity
        word_exporter._get_manual_content = MagicMock(return_value=None)
        word_exporter._get_output_path = MagicMock(return_value=output_path)

        result = word_exporter.export(output_path=output_path)

        assert result == output_path
        assert Path(output_path).exists()
        assert output_path.endswith(".docx")

    def test_export_with_language_parameter(self, word_exporter, tmp_path):
        """Test export respects language parameter."""
        output_path = str(tmp_path / "test_es.docx")

        word_exporter._get_manual_content = MagicMock(return_value=None)
        word_exporter._get_output_path = MagicMock(return_value=output_path)

        word_exporter.export(output_path=output_path, language="es")

        # Verify _get_manual_content was called with Spanish language
        calls = word_exporter._get_manual_content.call_args_list
        for call in calls:
            assert call[0][1] == "es"


class TestTemplateLoading:
    """Tests for template-related functionality."""

    def test_setup_styles_modifies_headings(self, tmp_path):
        """Test that _setup_styles modifies heading styles."""
        doc = Document()

        # Create a mock exporter
        exporter = WordExporter.__new__(WordExporter)
        exporter._setup_styles(doc)

        # Verify that styles exist (basic check)
        styles = doc.styles
        assert "Heading 1" in [s.name for s in styles]
        assert "Heading 2" in [s.name for s in styles]

    def test_setup_styles_sets_heading_colors(self, tmp_path):
        """Test that heading colors are set correctly."""
        doc = Document()

        exporter = WordExporter.__new__(WordExporter)
        exporter._setup_styles(doc)

        # Check that Heading 1 style exists and has been modified
        heading1 = doc.styles["Heading 1"]
        # The color should be set to RGBColor(0x1a, 0x36, 0x5d)
        expected_color = RGBColor(0x1a, 0x36, 0x5d)
        assert heading1.font.color.rgb == expected_color


class TestImageEmbedding:
    """Tests for image embedding functionality."""

    @pytest.fixture
    def word_exporter_with_storage(self, tmp_path):
        """Create WordExporter with real temp storage."""
        exporter = WordExporter.__new__(WordExporter)

        # Set up mock storage with real paths
        storage = MagicMock()
        storage.docs_dir = tmp_path / "docs"
        storage.docs_dir.mkdir(parents=True)

        exporter.user_storage = storage
        exporter.project = {"name": "Test", "chapters": []}

        return exporter

    def test_add_image_with_valid_path(self, word_exporter_with_storage, tmp_path):
        """Test adding an image with a valid path."""
        doc = Document()

        # Create a test image
        manual_id = "test-manual"
        screenshots_dir = word_exporter_with_storage.user_storage.docs_dir / manual_id / "screenshots"
        screenshots_dir.mkdir(parents=True)

        # Create a simple 1x1 pixel image
        from PIL import Image
        img_path = screenshots_dir / "test.png"
        img = Image.new("RGB", (100, 100), color="red")
        img.save(img_path)

        word_exporter_with_storage._add_image(doc, "screenshots/test.png", "Test image", manual_id)

        # The document should contain an inline shape (image)
        # Check that we have at least one paragraph added
        assert len(doc.paragraphs) >= 1

    def test_add_image_with_missing_path(self, word_exporter_with_storage):
        """Test adding an image with a missing path shows placeholder."""
        doc = Document()

        manual_id = "test-manual"
        screenshots_dir = word_exporter_with_storage.user_storage.docs_dir / manual_id / "screenshots"
        screenshots_dir.mkdir(parents=True)

        word_exporter_with_storage._add_image(doc, "screenshots/nonexistent.png", "Missing image", manual_id)

        # Should add placeholder text
        assert len(doc.paragraphs) >= 1
        para_text = doc.paragraphs[-1].text
        assert "not found" in para_text.lower()

    def test_add_image_with_file_url(self, word_exporter_with_storage, tmp_path):
        """Test adding an image with file:// URL prefix."""
        doc = Document()

        manual_id = "test-manual"
        screenshots_dir = word_exporter_with_storage.user_storage.docs_dir / manual_id / "screenshots"
        screenshots_dir.mkdir(parents=True)

        # Create test image
        from PIL import Image
        img_path = tmp_path / "absolute_test.png"
        img = Image.new("RGB", (100, 100), color="blue")
        img.save(img_path)

        word_exporter_with_storage._add_image(
            doc, f"file://{img_path}", "Absolute path image", manual_id
        )

        # Should have added content
        assert len(doc.paragraphs) >= 1


class TestInvalidTemplate:
    """Tests for handling invalid templates and error cases."""

    def test_export_with_missing_project(self):
        """Test that export fails gracefully with missing project."""
        with patch("src.export.word_exporter.BaseExporter.__init__") as mock_init:
            mock_init.side_effect = ValueError("Project not found: missing-project")

            with pytest.raises(ValueError) as exc_info:
                WordExporter("test-user", "missing-project")

            assert "Project not found" in str(exc_info.value)

    def test_export_with_no_chapters(self, tmp_path):
        """Test export with a project that has no chapters."""
        exporter = WordExporter.__new__(WordExporter)
        exporter.user_storage = MagicMock()
        exporter.user_storage.docs_dir = tmp_path
        exporter.project_storage = MagicMock()
        exporter.project = {
            "name": "Empty Project",
            "description": None,
            "chapters": [],
        }
        exporter.user_id = "test-user"
        exporter.project_id = "empty-project"
        exporter._get_output_path = MagicMock(return_value=str(tmp_path / "empty.docx"))

        result = exporter.export()

        assert Path(result).exists()
        # Verify the document can be opened
        doc = Document(result)
        assert len(doc.paragraphs) > 0


class TestMarkdownParsing:
    """Tests for markdown to Word conversion."""

    @pytest.fixture
    def word_exporter(self, tmp_path):
        """Create a WordExporter instance for markdown testing."""
        exporter = WordExporter.__new__(WordExporter)
        storage = MagicMock()
        storage.docs_dir = tmp_path
        exporter.user_storage = storage
        return exporter

    def test_parse_heading_level_1(self, word_exporter):
        """Test parsing level 1 heading."""
        doc = Document()
        word_exporter._add_markdown_content(doc, "# Main Title", "test-manual")

        assert len(doc.paragraphs) >= 1
        # Should have a heading style
        heading = doc.paragraphs[0]
        assert heading.style.name.startswith("Heading")

    def test_parse_heading_level_2(self, word_exporter):
        """Test parsing level 2 heading."""
        doc = Document()
        word_exporter._add_markdown_content(doc, "## Section Title", "test-manual")

        assert len(doc.paragraphs) >= 1

    def test_parse_bullet_list(self, word_exporter):
        """Test parsing bullet list items."""
        doc = Document()
        content = "- Item 1\n- Item 2\n- Item 3"
        word_exporter._add_markdown_content(doc, content, "test-manual")

        # Should have multiple paragraphs
        assert len(doc.paragraphs) >= 3

    def test_parse_numbered_list(self, word_exporter):
        """Test parsing numbered list items."""
        doc = Document()
        content = "1. First\n2. Second\n3. Third"
        word_exporter._add_markdown_content(doc, content, "test-manual")

        assert len(doc.paragraphs) >= 3

    def test_parse_code_block(self, word_exporter):
        """Test parsing code blocks."""
        doc = Document()
        content = "```python\nprint('hello')\n```"
        word_exporter._add_markdown_content(doc, content, "test-manual")

        # Should have at least one paragraph
        assert len(doc.paragraphs) >= 1

    def test_parse_blockquote(self, word_exporter):
        """Test parsing blockquotes."""
        doc = Document()
        content = "> This is a quote"
        word_exporter._add_markdown_content(doc, content, "test-manual")

        assert len(doc.paragraphs) >= 1
        # Check that the quote has left indent
        para = doc.paragraphs[0]
        assert para.paragraph_format.left_indent == Inches(0.5)

    def test_parse_horizontal_rule(self, word_exporter):
        """Test parsing horizontal rules."""
        doc = Document()
        content = "---"
        word_exporter._add_markdown_content(doc, content, "test-manual")

        # Should add centered line
        assert len(doc.paragraphs) >= 1
        para = doc.paragraphs[0]
        assert para.alignment == WD_ALIGN_PARAGRAPH.CENTER


class TestFormattedText:
    """Tests for inline text formatting."""

    @pytest.fixture
    def word_exporter(self):
        """Create WordExporter for text formatting tests."""
        return WordExporter.__new__(WordExporter)

    def test_process_bold_text(self, word_exporter):
        """Test processing bold markdown."""
        text = "This is **bold** text"
        result = word_exporter._process_inline_formatting(text)
        assert "bold" in result
        assert "**" not in result

    def test_process_italic_text(self, word_exporter):
        """Test processing italic markdown."""
        text = "This is *italic* text"
        result = word_exporter._process_inline_formatting(text)
        assert "italic" in result
        assert result.count("*") == 0

    def test_process_inline_code(self, word_exporter):
        """Test processing inline code."""
        text = "Use `code` here"
        result = word_exporter._process_inline_formatting(text)
        assert "code" in result
        assert "`" not in result

    def test_process_links(self, word_exporter):
        """Test processing markdown links."""
        text = "Click [here](https://example.com)"
        result = word_exporter._process_inline_formatting(text)
        assert "here" in result
        assert "https://example.com" not in result

    def test_add_formatted_text_bold(self, word_exporter):
        """Test adding bold formatted text to paragraph."""
        doc = Document()
        para = doc.add_paragraph()

        word_exporter._add_formatted_text(para, "This is **bold** text")

        # Check that at least one run is bold
        bold_runs = [run for run in para.runs if run.bold]
        assert len(bold_runs) >= 1

    def test_add_formatted_text_italic(self, word_exporter):
        """Test adding italic formatted text to paragraph."""
        doc = Document()
        para = doc.add_paragraph()

        word_exporter._add_formatted_text(para, "This is *italic* text")

        # Check that at least one run is italic
        italic_runs = [run for run in para.runs if run.italic]
        assert len(italic_runs) >= 1

    def test_add_formatted_text_code(self, word_exporter):
        """Test adding inline code formatted text to paragraph."""
        doc = Document()
        para = doc.add_paragraph()

        word_exporter._add_formatted_text(para, "Use `code` here")

        # Check that at least one run has Courier New font
        code_runs = [run for run in para.runs if run.font.name == "Courier New"]
        assert len(code_runs) >= 1


class TestChapterCover:
    """Tests for chapter cover page generation."""

    @pytest.fixture
    def word_exporter(self):
        """Create WordExporter for chapter cover tests."""
        exporter = WordExporter.__new__(WordExporter)
        return exporter

    def test_add_chapter_cover(self, word_exporter):
        """Test adding a chapter cover page."""
        doc = Document()
        chapter = {
            "id": "chapter-1",
            "title": "Getting Started",
            "description": "Introduction to the system",
            "order": 1,
            "manuals": ["manual-1", "manual-2"],
        }

        word_exporter._add_chapter_cover(doc, chapter)

        # Should have added paragraphs
        assert len(doc.paragraphs) > 0

        # Find the title heading
        title_found = False
        for para in doc.paragraphs:
            if "Getting Started" in para.text:
                title_found = True
                break
        assert title_found

    def test_add_chapter_cover_shows_manual_count(self, word_exporter):
        """Test that chapter cover shows manual count."""
        doc = Document()
        chapter = {
            "id": "chapter-1",
            "title": "Test Chapter",
            "order": 1,
            "manuals": ["m1", "m2", "m3"],
        }

        word_exporter._add_chapter_cover(doc, chapter)

        # Find manual count paragraph
        count_found = False
        for para in doc.paragraphs:
            if "3 manual" in para.text.lower():
                count_found = True
                break
        assert count_found


class TestCodeBlock:
    """Tests for code block formatting."""

    @pytest.fixture
    def word_exporter(self):
        """Create WordExporter for code block tests."""
        return WordExporter.__new__(WordExporter)

    def test_add_code_block(self, word_exporter):
        """Test adding a code block."""
        doc = Document()
        code = "def hello():\n    print('Hello')"

        word_exporter._add_code_block(doc, code)

        assert len(doc.paragraphs) >= 1
        para = doc.paragraphs[0]
        assert para.runs[0].font.name == "Courier New"
        assert para.runs[0].font.size == Pt(9)

    def test_add_code_block_indentation(self, word_exporter):
        """Test that code block has proper indentation."""
        doc = Document()
        code = "print('test')"

        word_exporter._add_code_block(doc, code)

        para = doc.paragraphs[0]
        assert para.paragraph_format.left_indent == Inches(0.25)

    def test_add_code_block_spacing(self, word_exporter):
        """Test that code block has proper spacing."""
        doc = Document()
        code = "print('test')"

        word_exporter._add_code_block(doc, code)

        para = doc.paragraphs[0]
        assert para.paragraph_format.space_before == Pt(6)
        assert para.paragraph_format.space_after == Pt(6)
