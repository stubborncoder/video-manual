"""Word (.docx) exporter for projects using python-docx."""

import base64
import re
from io import BytesIO
from pathlib import Path
from typing import Optional, List, Tuple

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from .base_exporter import BaseExporter


class WordExporter(BaseExporter):
    """Exports projects to Word (.docx) format."""

    file_extension = "docx"

    def export(
        self,
        output_path: Optional[str] = None,
        language: str = "en",
        include_toc: bool = False,
        include_chapter_covers: bool = False,
    ) -> str:
        """Export project to Word document.

        Args:
            output_path: Optional output file path
            language: Language code for manual content
            include_toc: Whether to include table of contents
            include_chapter_covers: Whether to include chapter cover pages

        Returns:
            Path to the generated Word file
        """
        output_path = self._get_output_path(output_path, language)

        # Create Word document
        doc = Document()
        self._setup_styles(doc)

        # Add title
        title = doc.add_heading(self.project['name'], level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if self.project.get('description'):
            desc = doc.add_paragraph(self.project['description'])
            desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            desc.runs[0].italic = True

        # Add page break after title
        doc.add_page_break()

        # Process chapters and manuals
        chapters = self.project.get("chapters", [])
        for chapter in sorted(chapters, key=lambda c: c.get("order", 0)):
            if include_chapter_covers:
                self._add_chapter_cover(doc, chapter)

            for manual_id in chapter.get("manuals", []):
                content = self._get_manual_content(manual_id, language)
                if content:
                    self._add_markdown_content(doc, content, manual_id)
                else:
                    doc.add_heading(manual_id, level=1)
                    doc.add_paragraph(f"Manual not found for language: {language}")

        doc.save(output_path)
        return output_path

    def _setup_styles(self, doc: Document) -> None:
        """Set up document styles."""
        styles = doc.styles

        # Modify existing heading styles
        for i in range(1, 5):
            style_name = f'Heading {i}'
            if style_name in styles:
                style = styles[style_name]
                style.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)

    def _add_chapter_cover(self, doc: Document, chapter: dict) -> None:
        """Add a chapter cover page."""
        # Add spacing
        for _ in range(10):
            doc.add_paragraph()

        # Chapter title
        title = doc.add_heading(f"Chapter {chapter.get('order', '')}: {chapter['title']}", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if chapter.get('description'):
            desc = doc.add_paragraph(chapter['description'])
            desc.alignment = WD_ALIGN_PARAGRAPH.CENTER

        manual_count = len(chapter.get('manuals', []))
        count_para = doc.add_paragraph(f"{manual_count} manual(s)")
        count_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

    def _add_markdown_content(self, doc: Document, content: str, manual_id: str) -> None:
        """Parse markdown content and add to Word document.

        Args:
            doc: Word document
            content: Markdown content
            manual_id: Manual identifier (for resolving image paths)
        """
        lines = content.split('\n')
        i = 0
        in_code_block = False
        code_content = []

        while i < len(lines):
            line = lines[i]

            # Handle code blocks
            if line.startswith('```'):
                if in_code_block:
                    # End code block
                    self._add_code_block(doc, '\n'.join(code_content))
                    code_content = []
                    in_code_block = False
                else:
                    # Start code block
                    in_code_block = True
                i += 1
                continue

            if in_code_block:
                code_content.append(line)
                i += 1
                continue

            # Handle headings
            if line.startswith('#'):
                level = len(line.split()[0]) if line.split() else 1
                text = line.lstrip('#').strip()
                if text:
                    doc.add_heading(text, level=min(level, 9))
                i += 1
                continue

            # Handle images
            img_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line)
            if img_match:
                alt_text = img_match.group(1)
                img_path = img_match.group(2)
                self._add_image(doc, img_path, alt_text, manual_id)
                i += 1
                continue

            # Handle blockquotes
            if line.startswith('>'):
                text = line.lstrip('>').strip()
                para = doc.add_paragraph(text)
                para.style = 'Quote' if 'Quote' in doc.styles else None
                para.paragraph_format.left_indent = Inches(0.5)
                i += 1
                continue

            # Handle horizontal rules
            if line.strip() in ['---', '***', '___']:
                # Add a thin horizontal line
                para = doc.add_paragraph('_' * 50)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                i += 1
                continue

            # Handle unordered lists
            if line.strip().startswith(('- ', '* ', '+ ')):
                text = line.strip()[2:]
                para = doc.add_paragraph(text, style='List Bullet')
                i += 1
                continue

            # Handle ordered lists
            ol_match = re.match(r'^\s*(\d+)\.\s+(.+)', line)
            if ol_match:
                text = ol_match.group(2)
                para = doc.add_paragraph(text, style='List Number')
                i += 1
                continue

            # Handle regular paragraphs
            if line.strip():
                # Process inline formatting
                text = self._process_inline_formatting(line)
                para = doc.add_paragraph()
                self._add_formatted_text(para, line)
            else:
                # Empty line - skip
                pass

            i += 1

    def _add_code_block(self, doc: Document, code: str) -> None:
        """Add a code block to the document."""
        para = doc.add_paragraph()
        run = para.add_run(code)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        para.paragraph_format.left_indent = Inches(0.25)
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)

    def _add_image(self, doc: Document, img_path: str, alt_text: str, manual_id: str) -> None:
        """Add an image to the document.

        Args:
            doc: Word document
            img_path: Image path (may be relative)
            alt_text: Alt text for the image
            manual_id: Manual identifier
        """
        # Resolve image path
        screenshots_dir = self.user_storage.manuals_dir / manual_id / "screenshots"

        if img_path.startswith("../screenshots/"):
            filename = img_path.replace("../screenshots/", "")
            abs_path = screenshots_dir / filename
        elif img_path.startswith("screenshots/"):
            filename = img_path.replace("screenshots/", "")
            abs_path = screenshots_dir / filename
        elif img_path.startswith("file://"):
            abs_path = Path(img_path.replace("file://", ""))
        else:
            abs_path = Path(img_path)

        if abs_path.exists():
            try:
                # Add image with max width of 6 inches
                doc.add_picture(str(abs_path), width=Inches(6))

                # Add caption if alt text provided
                if alt_text:
                    caption = doc.add_paragraph(alt_text)
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption.runs[0].italic = True
                    caption.runs[0].font.size = Pt(9)
            except Exception:
                # If image fails to load, add placeholder text
                doc.add_paragraph(f"[Image: {alt_text or img_path}]")
        else:
            doc.add_paragraph(f"[Image not found: {img_path}]")

    def _process_inline_formatting(self, text: str) -> str:
        """Process inline markdown formatting (for display purposes)."""
        # Remove markdown formatting for plain text
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # Bold
        text = re.sub(r'\*(.+?)\*', r'\1', text)  # Italic
        text = re.sub(r'`(.+?)`', r'\1', text)  # Code
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)  # Links
        return text

    def _add_formatted_text(self, para, text: str) -> None:
        """Add text with inline formatting to a paragraph.

        Args:
            para: Paragraph to add text to
            text: Markdown text with potential formatting
        """
        # Simple approach: parse and add formatted runs
        # This handles **bold**, *italic*, `code`, and [links](url)

        pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\)|[^*`\[]+)'
        parts = re.findall(pattern, text)

        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold
                run = para.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                # Italic
                run = para.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                # Inline code
                run = para.add_run(part[1:-1])
                run.font.name = 'Courier New'
            elif part.startswith('['):
                # Link - extract text only
                match = re.match(r'\[([^\]]+)\]\(([^)]+)\)', part)
                if match:
                    run = para.add_run(match.group(1))
                    run.underline = True
                    run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
            else:
                # Regular text
                if part.strip():
                    para.add_run(part)
