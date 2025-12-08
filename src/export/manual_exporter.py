"""Manual exporter for exporting individual manuals to various formats."""

import re
from abc import ABC, abstractmethod
from datetime import datetime
from pathlib import Path
from typing import Optional
import base64

import markdown
from weasyprint import HTML, CSS
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from ..storage.user_storage import UserStorage
from .tag_parser import strip_semantic_tags


def slugify(text: str) -> str:
    """Convert text to URL-friendly slug."""
    text = text.lower()
    text = re.sub(r'[\s_]+', '-', text)
    text = re.sub(r'[^a-z0-9\-]', '', text)
    text = re.sub(r'-+', '-', text)
    return text.strip('-')


class BaseManualExporter(ABC):
    """Abstract base class for manual exporters."""

    # File extension for this export format (e.g., 'pdf', 'docx', 'html')
    file_extension: str = ""

    def __init__(self, user_id: str, manual_id: str):
        """Initialize exporter.

        Args:
            user_id: User identifier
            manual_id: Manual identifier
        """
        self.user_id = user_id
        self.manual_id = manual_id
        self.user_storage = UserStorage(user_id)

        # Verify manual exists
        manual_dir = self.user_storage.manuals_dir / manual_id
        if not manual_dir.exists():
            raise ValueError(f"Manual not found: {manual_id}")

    @abstractmethod
    def export(
        self,
        output_path: Optional[str] = None,
        language: str = "en",
        **options,
    ) -> str:
        """Export manual to the target format.

        Args:
            output_path: Optional output file path
            language: Language code for manual content
            **options: Format-specific options

        Returns:
            Path to the generated file
        """
        pass

    def _get_output_path(
        self,
        output_path: Optional[str],
        language: str,
    ) -> str:
        """Get or generate output path for export.

        Args:
            output_path: Optional explicit output path
            language: Language code

        Returns:
            Output file path
        """
        if output_path:
            return output_path

        export_dir = self.user_storage.manuals_dir / self.manual_id / "exports"
        export_dir.mkdir(parents=True, exist_ok=True)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{self.manual_id}_{language}_{timestamp}.{self.file_extension}"
        return str(export_dir / filename)

    def _get_manual_content(self, language: str) -> Optional[str]:
        """Get manual content for a specific language.

        Strips semantic tags from the content so that exported documents
        contain clean markdown without XML-like tags.
        """
        content = self.user_storage.get_manual_content(self.manual_id, language)
        if content:
            # Strip semantic tags (e.g., <step>, <note>, <introduction>) for clean export
            content = strip_semantic_tags(content)
        return content

    def _fix_image_paths(self, content: str) -> str:
        """Fix relative image paths to absolute paths.

        Args:
            content: Markdown content

        Returns:
            Content with fixed image paths
        """
        screenshots_dir = self.user_storage.manuals_dir / self.manual_id / "screenshots"

        def replace_path(match):
            # Normalize alt text (may span multiple lines)
            alt = match.group(1).replace('\n', ' ').strip()
            alt = re.sub(r'\s+', ' ', alt)
            path = match.group(2)

            # Handle ../screenshots/ paths
            if path.startswith("../screenshots/"):
                filename = path.replace("../screenshots/", "")
                abs_path = screenshots_dir / filename
                if abs_path.exists():
                    return f"![{alt}](file://{abs_path})"

            # Handle screenshots/ paths
            if path.startswith("screenshots/"):
                filename = path.replace("screenshots/", "")
                abs_path = screenshots_dir / filename
                if abs_path.exists():
                    return f"![{alt}](file://{abs_path})"

            # Handle bare filenames
            if not path.startswith(("http://", "https://", "file://", "/")):
                abs_path = screenshots_dir / path
                if abs_path.exists():
                    return f"![{alt}](file://{abs_path})"

            return f"![{alt}]({path})"

        # Use [\s\S]*? to match alt text that may span multiple lines
        pattern = r'!\[([\s\S]*?)\]\(([^)]+)\)'
        return re.sub(pattern, replace_path, content)


class ManualPDFExporter(BaseManualExporter):
    """Export manual to PDF format."""

    file_extension = "pdf"

    # Default CSS for PDF export
    DEFAULT_CSS = """
@page {
    size: A4;
    margin: 2cm;
    @bottom-center {
        content: "Page " counter(page) " of " counter(pages);
        font-size: 10pt;
        color: #666;
    }
}

body {
    font-family: "Helvetica Neue", Helvetica, Arial, sans-serif;
    font-size: 11pt;
    line-height: 1.6;
    color: #333;
}

h1 {
    font-size: 24pt;
    color: #1a365d;
    border-bottom: 2px solid #1a365d;
    padding-bottom: 10px;
    margin-top: 0;
}

h2 {
    font-size: 18pt;
    color: #2c5282;
    margin-top: 30px;
    border-bottom: 1px solid #e2e8f0;
    padding-bottom: 5px;
    page-break-before: always;
}

h3 {
    font-size: 14pt;
    color: #2d3748;
    margin-top: 25px;
}

h4 {
    font-size: 12pt;
    color: #4a5568;
}

p {
    margin: 10px 0;
}

img {
    max-width: 100%;
    max-height: 400px;
    width: auto;
    height: auto;
    object-fit: contain;
    display: block;
    margin: 15px auto;
    border: 1px solid #e2e8f0;
    border-radius: 4px;
}

code {
    background-color: #f7fafc;
    padding: 2px 6px;
    border-radius: 3px;
    font-family: "Monaco", "Consolas", monospace;
    font-size: 10pt;
}

pre {
    background-color: #f7fafc;
    padding: 15px;
    border-radius: 5px;
    overflow-x: auto;
    border: 1px solid #e2e8f0;
}

pre code {
    padding: 0;
    background: none;
}

ul, ol {
    margin: 10px 0;
    padding-left: 30px;
}

li {
    margin: 5px 0;
}

table {
    width: 100%;
    border-collapse: collapse;
    margin: 15px 0;
}

th, td {
    border: 1px solid #e2e8f0;
    padding: 8px;
    text-align: left;
}

th {
    background-color: #f7fafc;
    font-weight: bold;
}

blockquote {
    border-left: 4px solid #cbd5e0;
    padding-left: 15px;
    margin: 15px 0;
    color: #4a5568;
    font-style: italic;
}

.manual-header {
    text-align: center;
    margin-bottom: 30px;
    padding-bottom: 20px;
    border-bottom: 2px solid #1a365d;
}

.manual-header h1 {
    border-bottom: none;
}

.manual-header .meta {
    font-size: 10pt;
    color: #718096;
    margin-top: 10px;
}
"""

    def export(
        self,
        output_path: Optional[str] = None,
        language: str = "en",
        **options,
    ) -> str:
        """Export manual to PDF.

        Args:
            output_path: Optional output file path
            language: Language code
            **options: Additional options (currently unused)

        Returns:
            Path to generated PDF file
        """
        # Get manual content
        content = self._get_manual_content(language)
        if not content:
            raise ValueError(f"Manual content not found for language: {language}")

        # Get manual metadata for title
        metadata = self.user_storage.get_manual_metadata(self.manual_id) or {}
        title = metadata.get("title") or self.manual_id

        # Fix image paths
        content = self._fix_image_paths(content)

        # Get language display name
        language_names = {
            "en": "English", "es": "Spanish", "fr": "French", "de": "German",
            "it": "Italian", "pt": "Portuguese", "nl": "Dutch", "pl": "Polish",
            "ru": "Russian", "zh": "Chinese", "ja": "Japanese", "ko": "Korean",
        }
        language_display = language_names.get(language, language.upper())

        # Add header
        header = f"""
<div class="manual-header">
    <h1>{title}</h1>
    <div class="meta">
        Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')} | Language: {language_display}
    </div>
</div>
"""

        # Convert markdown to HTML
        md = markdown.Markdown(extensions=['extra', 'codehilite', 'tables', 'toc'])
        html_content = md.convert(content)

        # Wrap in HTML template
        html = f"""
<!DOCTYPE html>
<html lang="{language}">
<head>
    <meta charset="UTF-8">
    <title>{title}</title>
</head>
<body>
{header}
{html_content}
</body>
</html>
"""

        # Generate PDF
        output_path = self._get_output_path(output_path, language)
        HTML(string=html).write_pdf(
            output_path,
            stylesheets=[CSS(string=self.DEFAULT_CSS)]
        )

        return output_path


class ManualWordExporter(BaseManualExporter):
    """Export manual to Word (DOCX) format."""

    file_extension = "docx"

    def export(
        self,
        output_path: Optional[str] = None,
        language: str = "en",
        **options,
    ) -> str:
        """Export manual to Word document.

        Args:
            output_path: Optional output file path
            language: Language code
            **options: Additional options (currently unused)

        Returns:
            Path to generated DOCX file
        """
        # Get manual content
        content = self._get_manual_content(language)
        if not content:
            raise ValueError(f"Manual content not found for language: {language}")

        # Create document
        doc = Document()

        # Add title
        title = doc.add_heading(self.manual_id, 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add metadata
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        meta_run = meta_para.add_run(
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')} | Language: {language.upper()}"
        )
        meta_run.font.size = Pt(10)
        meta_run.font.color.rgb = RGBColor(113, 128, 150)

        doc.add_paragraph()  # Spacing

        # Convert markdown to HTML first, then parse
        md = markdown.Markdown(extensions=['extra', 'codehilite', 'tables', 'toc'])
        html_content = md.convert(content)

        # Simple markdown parsing (basic implementation)
        # For a more robust solution, consider using a markdown parser
        self._add_markdown_to_doc(doc, content)

        # Save document
        output_path = self._get_output_path(output_path, language)
        doc.save(output_path)

        return output_path

    def _add_markdown_to_doc(self, doc: Document, content: str):
        """Add markdown content to Word document.

        This is a simplified markdown parser. For production use,
        consider using a more robust markdown-to-docx library.
        """
        screenshots_dir = self.user_storage.manuals_dir / self.manual_id / "screenshots"

        lines = content.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i]

            # Headers
            if line.startswith('# '):
                doc.add_heading(line[2:], 1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], 2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], 3)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], 4)
            # Images
            elif line.startswith('!['):
                match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line)
                if match:
                    alt_text = match.group(1)
                    img_path = match.group(2)

                    # Resolve image path
                    if img_path.startswith("../screenshots/"):
                        img_path = screenshots_dir / img_path.replace("../screenshots/", "")
                    elif img_path.startswith("screenshots/"):
                        img_path = screenshots_dir / img_path.replace("screenshots/", "")
                    elif not img_path.startswith(("http://", "https://", "file://", "/")):
                        img_path = screenshots_dir / img_path

                    # Add image if it exists
                    if isinstance(img_path, Path) and img_path.exists():
                        try:
                            # Check image dimensions to handle portrait vs landscape
                            from PIL import Image as PILImage
                            with PILImage.open(img_path) as pil_img:
                                img_width, img_height = pil_img.size

                            is_portrait = img_height > img_width
                            if is_portrait:
                                # Portrait: limit height to ~4 inches, scale width proportionally
                                max_height = Inches(4)
                                aspect_ratio = img_width / img_height
                                width = Inches(4 * aspect_ratio)
                                doc.add_picture(str(img_path), width=width)
                            else:
                                # Landscape: use full width
                                doc.add_picture(str(img_path), width=Inches(6))

                            if alt_text:
                                caption = doc.add_paragraph(alt_text)
                                caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                caption.runs[0].font.size = Pt(9)
                                caption.runs[0].font.color.rgb = RGBColor(113, 128, 150)
                        except Exception:
                            # If image fails to load, add alt text as paragraph
                            doc.add_paragraph(f"[Image: {alt_text}]")
            # Lists
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                doc.add_paragraph(line.strip()[2:], style='List Bullet')
            elif re.match(r'^\d+\.\s', line.strip()):
                text = re.sub(r'^\d+\.\s', '', line.strip())
                doc.add_paragraph(text, style='List Number')
            # Code blocks
            elif line.strip().startswith('```'):
                # Collect code block
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                if code_lines:
                    code_para = doc.add_paragraph()
                    code_run = code_para.add_run('\n'.join(code_lines))
                    code_run.font.name = 'Consolas'
                    code_run.font.size = Pt(9)
                    code_para.paragraph_format.left_indent = Inches(0.25)
                    code_para.paragraph_format.space_before = Pt(6)
                    code_para.paragraph_format.space_after = Pt(6)
            # Regular paragraphs
            elif line.strip():
                doc.add_paragraph(line)

            i += 1


class ManualHTMLExporter(BaseManualExporter):
    """Export manual to standalone HTML format."""

    file_extension = "html"

    # Default CSS for HTML export
    DEFAULT_CSS = """
body {
    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif;
    font-size: 16px;
    line-height: 1.6;
    color: #24292e;
    max-width: 900px;
    margin: 0 auto;
    padding: 20px;
    background-color: #ffffff;
}

@media (prefers-color-scheme: dark) {
    body {
        background-color: #0d1117;
        color: #c9d1d9;
    }

    h1, h2, h3, h4, h5, h6 {
        color: #c9d1d9;
    }

    code {
        background-color: #161b22;
        color: #c9d1d9;
    }

    pre {
        background-color: #161b22;
        border-color: #30363d;
    }

    blockquote {
        border-color: #3b434b;
        color: #8b949e;
    }

    table th, table td {
        border-color: #30363d;
    }

    table th {
        background-color: #161b22;
    }
}

.manual-header {
    text-align: center;
    margin-bottom: 40px;
    padding-bottom: 20px;
    border-bottom: 2px solid #e1e4e8;
}

.manual-header h1 {
    margin: 0 0 10px 0;
    font-size: 2.5em;
    color: #0366d6;
}

.manual-header .meta {
    font-size: 0.9em;
    color: #586069;
}

h1, h2, h3, h4, h5, h6 {
    margin-top: 24px;
    margin-bottom: 16px;
    font-weight: 600;
    line-height: 1.25;
}

h1 { font-size: 2em; border-bottom: 1px solid #e1e4e8; padding-bottom: 0.3em; }
h2 { font-size: 1.5em; border-bottom: 1px solid #e1e4e8; padding-bottom: 0.3em; }
h3 { font-size: 1.25em; }
h4 { font-size: 1em; }

p {
    margin-top: 0;
    margin-bottom: 16px;
}

img {
    max-width: 100%;
    height: auto;
    display: block;
    margin: 20px auto;
    border: 1px solid #e1e4e8;
    border-radius: 6px;
    box-shadow: 0 1px 3px rgba(0,0,0,0.12);
}

code {
    background-color: #f6f8fa;
    padding: 0.2em 0.4em;
    border-radius: 3px;
    font-family: "SFMono-Regular", Consolas, "Liberation Mono", Menlo, monospace;
    font-size: 85%;
}

pre {
    background-color: #f6f8fa;
    padding: 16px;
    border-radius: 6px;
    overflow: auto;
    border: 1px solid #e1e4e8;
}

pre code {
    background: none;
    padding: 0;
}

ul, ol {
    padding-left: 2em;
    margin-top: 0;
    margin-bottom: 16px;
}

li {
    margin-top: 0.25em;
}

table {
    width: 100%;
    border-collapse: collapse;
    margin: 16px 0;
}

table th,
table td {
    padding: 6px 13px;
    border: 1px solid #dfe2e5;
}

table th {
    font-weight: 600;
    background-color: #f6f8fa;
}

blockquote {
    margin: 0;
    padding: 0 1em;
    color: #6a737d;
    border-left: 0.25em solid #dfe2e5;
}

a {
    color: #0366d6;
    text-decoration: none;
}

a:hover {
    text-decoration: underline;
}
"""

    def export(
        self,
        output_path: Optional[str] = None,
        language: str = "en",
        embed_images: bool = True,
        **options,
    ) -> str:
        """Export manual to standalone HTML.

        Args:
            output_path: Optional output file path
            language: Language code
            embed_images: Embed images as base64 data URIs
            **options: Additional options (currently unused)

        Returns:
            Path to generated HTML file
        """
        # Get manual content
        content = self._get_manual_content(language)
        if not content:
            raise ValueError(f"Manual content not found for language: {language}")

        # Embed images if requested
        if embed_images:
            content = self._embed_images_as_base64(content)
        else:
            content = self._fix_image_paths(content)

        # Add header
        header = f"""
<div class="manual-header">
    <h1>{self.manual_id}</h1>
    <div class="meta">
        Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')} | Language: {language.upper()}
    </div>
</div>
"""

        # Convert markdown to HTML
        md = markdown.Markdown(extensions=['extra', 'codehilite', 'tables', 'toc'])
        html_content = md.convert(content)

        # Wrap in HTML template
        html = f"""<!DOCTYPE html>
<html lang="{language}">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{self.manual_id}</title>
    <style>
{self.DEFAULT_CSS}
    </style>
</head>
<body>
{header}
{html_content}
</body>
</html>
"""

        # Save HTML
        output_path = self._get_output_path(output_path, language)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html)

        return output_path

    def _embed_images_as_base64(self, content: str) -> str:
        """Embed images as base64 data URIs.

        Args:
            content: Markdown content

        Returns:
            Content with embedded images
        """
        screenshots_dir = self.user_storage.manuals_dir / self.manual_id / "screenshots"

        def replace_path(match):
            alt = match.group(1).replace('\n', ' ').strip()
            alt = re.sub(r'\s+', ' ', alt)
            path = match.group(2)

            # Skip external URLs
            if path.startswith(("http://", "https://", "data:")):
                return f"![{alt}]({path})"

            # Resolve local path
            img_path = None
            if path.startswith("../screenshots/"):
                img_path = screenshots_dir / path.replace("../screenshots/", "")
            elif path.startswith("screenshots/"):
                img_path = screenshots_dir / path.replace("screenshots/", "")
            elif not path.startswith(("file://", "/")):
                img_path = screenshots_dir / path

            # Embed as base64 if file exists
            if img_path and img_path.exists():
                try:
                    with open(img_path, 'rb') as f:
                        img_data = f.read()
                    b64_data = base64.b64encode(img_data).decode('utf-8')

                    # Determine MIME type from extension
                    ext = img_path.suffix.lower()
                    mime_types = {
                        '.png': 'image/png',
                        '.jpg': 'image/jpeg',
                        '.jpeg': 'image/jpeg',
                        '.gif': 'image/gif',
                        '.svg': 'image/svg+xml',
                        '.webp': 'image/webp',
                    }
                    mime_type = mime_types.get(ext, 'image/png')

                    return f"![{alt}](data:{mime_type};base64,{b64_data})"
                except Exception:
                    pass

            return f"![{alt}]({path})"

        pattern = r'!\[([\s\S]*?)\]\(([^)]+)\)'
        return re.sub(pattern, replace_path, content)


def create_manual_exporter(user_id: str, manual_id: str, format: str) -> BaseManualExporter:
    """Factory function to create appropriate exporter.

    Args:
        user_id: User identifier
        manual_id: Manual identifier
        format: Export format ('pdf', 'word', 'html')

    Returns:
        Exporter instance

    Raises:
        ValueError: If format is not supported
    """
    exporters = {
        'pdf': ManualPDFExporter,
        'word': ManualWordExporter,
        'docx': ManualWordExporter,
        'html': ManualHTMLExporter,
    }

    exporter_class = exporters.get(format.lower())
    if not exporter_class:
        raise ValueError(f"Unsupported export format: {format}")

    return exporter_class(user_id, manual_id)
