"""Export compilation version markdown to various formats."""

import re
from datetime import datetime
from pathlib import Path
from typing import Optional

import markdown


def export_compilation_markdown(
    content: str,
    format: str,
    project_name: str,
    project_id: str,
    version: str,
    language: str,
    screenshots_dir: Optional[Path],
    output_dir: Path,
    user_id: str,
) -> Path:
    """Export pre-built compilation markdown to various formats.

    Args:
        content: Markdown content to export
        format: Export format (pdf, word, html)
        project_name: Project name for title
        project_id: Project ID
        version: Compilation version
        language: Content language
        screenshots_dir: Directory containing screenshots
        output_dir: Output directory for exports
        user_id: User ID

    Returns:
        Path to generated file
    """
    # Ensure output directory exists
    output_dir.mkdir(parents=True, exist_ok=True)

    # Generate output filename
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    # Fix image paths in content
    if screenshots_dir and screenshots_dir.exists():
        content = _fix_image_paths(content, screenshots_dir)

    if format == "pdf":
        return _export_pdf(
            content, project_name, output_dir,
            f"{project_id}_{language}_v{version}_{timestamp}.pdf",
            screenshots_dir,
        )
    elif format in ("word", "docx"):
        return _export_word(
            content, project_name, output_dir,
            f"{project_id}_{language}_v{version}_{timestamp}.docx",
            screenshots_dir,
        )
    elif format == "html":
        return _export_html(
            content, project_name, output_dir,
            f"{project_id}_{language}_v{version}_{timestamp}.html",
            screenshots_dir,
        )
    else:
        raise ValueError(f"Unknown format '{format}'. Use: pdf, word, html")


def _fix_image_paths(content: str, screenshots_dir: Path) -> str:
    """Fix relative image paths to absolute paths."""

    def replace_path(match):
        alt = match.group(1).replace('\n', ' ').strip()
        alt = re.sub(r'\s+', ' ', alt)
        path = match.group(2)

        # Handle screenshots/ paths
        if path.startswith("screenshots/"):
            filename = path.replace("screenshots/", "")
            abs_path = screenshots_dir / filename
            if abs_path.exists():
                return f"![{alt}](file://{abs_path})"

        # Handle direct filenames
        if not path.startswith(("http", "file://", "/")):
            abs_path = screenshots_dir / path
            if abs_path.exists():
                return f"![{alt}](file://{abs_path})"

        return f"![{alt}]({path})"

    pattern = r'!\[([\s\S]*?)\]\(([^)]+)\)'
    return re.sub(pattern, replace_path, content)


def _export_pdf(
    content: str,
    project_name: str,
    output_dir: Path,
    filename: str,
    screenshots_dir: Optional[Path],
) -> Path:
    """Export to PDF using WeasyPrint."""
    from weasyprint import HTML, CSS

    html_content = _convert_to_html(content, project_name)

    css = CSS(string=PDF_CSS)
    base_url = str(screenshots_dir) if screenshots_dir else None
    html_doc = HTML(string=html_content, base_url=base_url)

    output_path = output_dir / filename
    html_doc.write_pdf(str(output_path), stylesheets=[css])

    return output_path


def _export_word(
    content: str,
    project_name: str,
    output_dir: Path,
    filename: str,
    screenshots_dir: Optional[Path],
) -> Path:
    """Export to Word document."""
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title = doc.add_heading(project_name, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Process content line by line
    lines = content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        # Headers
        if line.startswith('######'):
            doc.add_heading(line[6:].strip(), level=6)
        elif line.startswith('#####'):
            doc.add_heading(line[5:].strip(), level=5)
        elif line.startswith('####'):
            doc.add_heading(line[4:].strip(), level=4)
        elif line.startswith('###'):
            doc.add_heading(line[3:].strip(), level=3)
        elif line.startswith('##'):
            doc.add_heading(line[2:].strip(), level=2)
        elif line.startswith('#'):
            doc.add_heading(line[1:].strip(), level=1)

        # Images
        elif line.startswith('!['):
            match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
            if match:
                alt_text, img_path = match.groups()
                img_path = img_path.replace('file://', '')

                if Path(img_path).exists():
                    try:
                        doc.add_picture(img_path, width=Inches(6))
                        if alt_text:
                            caption = doc.add_paragraph(alt_text)
                            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            caption.runs[0].italic = True
                    except Exception:
                        doc.add_paragraph(f"[Image: {alt_text}]")
                else:
                    doc.add_paragraph(f"[Image not found: {img_path}]")

        # Horizontal rule
        elif line.startswith('---') or line.startswith('***'):
            doc.add_paragraph('─' * 50)

        # List items
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\.\s', line):
            doc.add_paragraph(re.sub(r'^\d+\.\s', '', line), style='List Number')

        # Blockquote
        elif line.startswith('>'):
            p = doc.add_paragraph(line[1:].strip())
            p.paragraph_format.left_indent = Inches(0.5)
            p.runs[0].italic = True

        # Regular paragraph
        else:
            p = doc.add_paragraph()
            _add_formatted_text(p, line)

        i += 1

    output_path = output_dir / filename
    doc.save(str(output_path))

    return output_path


def _add_formatted_text(paragraph, text: str):
    """Add text with basic formatting (bold, italic, code)."""
    # Simple text for now - could be enhanced with regex for **bold**, *italic*, `code`
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
        else:
            paragraph.add_run(part)


def _export_html(
    content: str,
    project_name: str,
    output_dir: Path,
    filename: str,
    screenshots_dir: Optional[Path],
) -> Path:
    """Export to standalone HTML with embedded images."""
    import base64

    # Convert to HTML
    html_body = markdown.markdown(
        content,
        extensions=['tables', 'fenced_code', 'codehilite', 'toc', 'nl2br'],
    )

    # Embed images as base64
    if screenshots_dir:
        html_body = _embed_images_base64(html_body, screenshots_dir)

    html_doc = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{project_name}</title>
    <style>{HTML_CSS}</style>
</head>
<body>
    <article class="content">
        {html_body}
    </article>
</body>
</html>"""

    output_path = output_dir / filename
    output_path.write_text(html_doc, encoding='utf-8')

    return output_path


def _embed_images_base64(html: str, screenshots_dir: Path) -> str:
    """Embed images as base64 data URIs."""
    import base64

    def replace_src(match):
        src = match.group(1)
        src_clean = src.replace('file://', '')

        # Try to find the image
        img_path = None
        if Path(src_clean).exists():
            img_path = Path(src_clean)
        elif screenshots_dir:
            # Try relative path
            filename = Path(src_clean).name
            potential_path = screenshots_dir / filename
            if potential_path.exists():
                img_path = potential_path

        if img_path and img_path.exists():
            try:
                suffix = img_path.suffix.lower()
                mime_type = {
                    '.png': 'image/png',
                    '.jpg': 'image/jpeg',
                    '.jpeg': 'image/jpeg',
                    '.gif': 'image/gif',
                    '.webp': 'image/webp',
                }.get(suffix, 'image/png')

                with open(img_path, 'rb') as f:
                    data = base64.b64encode(f.read()).decode('utf-8')
                return f'src="data:{mime_type};base64,{data}"'
            except Exception:
                pass

        return match.group(0)

    return re.sub(r'src="([^"]+)"', replace_src, html)


def _convert_to_html(content: str, project_name: str) -> str:
    """Convert markdown to HTML document."""
    extensions = ['tables', 'fenced_code', 'codehilite', 'toc', 'nl2br', 'md_in_html']

    html_body = markdown.markdown(
        content,
        extensions=extensions,
        extension_configs={
            'codehilite': {'css_class': 'highlight', 'guess_lang': False}
        }
    )

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <title>{project_name}</title>
</head>
<body>
{html_body}
</body>
</html>"""


# CSS styles for PDF export
PDF_CSS = """
@page {
    size: A4;
    margin: 2cm;
    @bottom-center {
        content: "Page " counter(page) " of " counter(pages);
        font-size: 10pt;
        color: #666;
    }
}

body {
    font-family: "Helvetica Neue", Helvetica, Arial, sans-serif;
    font-size: 11pt;
    line-height: 1.6;
    color: #333;
}

h1 { font-size: 24pt; color: #1a365d; border-bottom: 2px solid #1a365d; padding-bottom: 10px; }
h2 { font-size: 18pt; color: #2c5282; margin-top: 30px; border-bottom: 1px solid #e2e8f0; padding-bottom: 5px; }
h3 { font-size: 14pt; color: #2d3748; margin-top: 25px; }
h4 { font-size: 12pt; color: #4a5568; }

img {
    max-width: 100%;
    height: auto;
    display: block;
    margin: 15px auto;
    border: 1px solid #e2e8f0;
    border-radius: 4px;
}

code {
    background-color: #f7fafc;
    padding: 2px 6px;
    border-radius: 3px;
    font-family: "Monaco", "Consolas", monospace;
    font-size: 10pt;
}

pre {
    background-color: #f7fafc;
    padding: 15px;
    border-radius: 5px;
    overflow-x: auto;
    border: 1px solid #e2e8f0;
}

blockquote {
    border-left: 4px solid #4299e1;
    margin: 15px 0;
    padding: 10px 20px;
    background-color: #ebf8ff;
    font-style: italic;
}

table { width: 100%; border-collapse: collapse; margin: 15px 0; }
th, td { border: 1px solid #e2e8f0; padding: 8px 12px; text-align: left; }
th { background-color: #f7fafc; font-weight: bold; }

hr { border: none; border-top: 1px solid #e2e8f0; margin: 30px 0; }
"""

# CSS styles for HTML export
HTML_CSS = """
body {
    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif;
    line-height: 1.6;
    max-width: 900px;
    margin: 0 auto;
    padding: 2rem;
    color: #333;
    background: #fff;
}

h1 { color: #1a365d; border-bottom: 2px solid #1a365d; padding-bottom: 0.5rem; }
h2 { color: #2c5282; margin-top: 2rem; border-bottom: 1px solid #e2e8f0; padding-bottom: 0.25rem; }
h3 { color: #2d3748; margin-top: 1.5rem; }

img {
    max-width: 100%;
    height: auto;
    display: block;
    margin: 1rem auto;
    border: 1px solid #e2e8f0;
    border-radius: 8px;
    box-shadow: 0 2px 8px rgba(0,0,0,0.1);
}

code {
    background-color: #f7fafc;
    padding: 0.2em 0.4em;
    border-radius: 3px;
    font-family: "Monaco", "Consolas", monospace;
    font-size: 0.9em;
}

pre {
    background-color: #f7fafc;
    padding: 1rem;
    border-radius: 8px;
    overflow-x: auto;
    border: 1px solid #e2e8f0;
}

blockquote {
    border-left: 4px solid #4299e1;
    margin: 1rem 0;
    padding: 0.5rem 1rem;
    background-color: #ebf8ff;
    border-radius: 0 8px 8px 0;
}

table { width: 100%; border-collapse: collapse; margin: 1rem 0; }
th, td { border: 1px solid #e2e8f0; padding: 0.5rem 0.75rem; text-align: left; }
th { background-color: #f7fafc; }

hr { border: none; border-top: 2px solid #e2e8f0; margin: 2rem 0; }

@media (prefers-color-scheme: dark) {
    body { background: #1a202c; color: #e2e8f0; }
    h1, h2, h3 { color: #90cdf4; }
    code, pre { background-color: #2d3748; }
    blockquote { background-color: #2d3748; }
    table, th, td { border-color: #4a5568; }
    th { background-color: #2d3748; }
}
"""
