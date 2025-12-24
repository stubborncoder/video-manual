#!/usr/bin/env python3
"""Create default Word templates for vDocs.

Run this script to generate the default templates in data/templates/.
These templates use Jinja2 syntax compatible with docxtpl.

Creates templates for each document format:
- step-manual: Step-by-step procedural manuals
- quick-guide: Brief overview with key points
- reference: Detailed reference documentation
- summary: Executive summary with findings

Usage:
    uv run python scripts/create_default_templates.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Twips, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Logo path for header
LOGO_PATH = Path(__file__).parent.parent / "data" / "templates" / "vdocs-logo.png"


def setup_base_styles(doc: Document) -> None:
    """Set up common styles used across all templates with modern typography."""
    styles = doc.styles

    # Modern font choices
    HEADING_FONT = "Segoe UI"  # Clean, modern sans-serif
    BODY_FONT = "Segoe UI"     # Consistent look

    # Normal/body text style
    normal_style = styles['Normal']
    normal_style.font.name = BODY_FONT
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.widow_control = True

    # Title style - Bold and impactful
    title_style = styles['Title']
    title_style.font.name = HEADING_FONT
    title_style.font.size = Pt(28)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    title_style.paragraph_format.space_after = Pt(8)
    title_style.paragraph_format.keep_with_next = True

    # Heading 1 - Section headers
    h1_style = styles['Heading 1']
    h1_style.font.name = HEADING_FONT
    h1_style.font.size = Pt(18)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)  # Primary blue
    h1_style.paragraph_format.space_before = Pt(16)
    h1_style.paragraph_format.space_after = Pt(6)
    h1_style.paragraph_format.keep_with_next = True

    # Heading 2 - Subsection headers (steps, items)
    h2_style = styles['Heading 2']
    h2_style.font.name = HEADING_FONT
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
    h2_style.paragraph_format.space_before = Pt(12)
    h2_style.paragraph_format.space_after = Pt(4)
    h2_style.paragraph_format.keep_with_next = True

    # Heading 3 - Minor headers
    h3_style = styles['Heading 3']
    h3_style.font.name = HEADING_FONT
    h3_style.font.size = Pt(12)
    h3_style.font.bold = True
    h3_style.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
    h3_style.paragraph_format.space_before = Pt(8)
    h3_style.paragraph_format.space_after = Pt(3)
    h3_style.paragraph_format.keep_with_next = True

    # Jinja control style - zero spacing for control tags
    try:
        styles['JinjaControl']
    except KeyError:
        jinja_style = styles.add_style('JinjaControl', WD_STYLE_TYPE.PARAGRAPH)
        jinja_style.font.size = Pt(1)  # Tiny font
        jinja_style.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)  # White (invisible)
        jinja_style.paragraph_format.space_before = Pt(0)
        jinja_style.paragraph_format.space_after = Pt(0)
        jinja_style.paragraph_format.line_spacing = Pt(1)

    # Create Code style for code blocks
    try:
        styles['Code']
    except KeyError:
        code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = 'Consolas'
        code_style.font.size = Pt(9)
        code_style.paragraph_format.space_before = Pt(4)
        code_style.paragraph_format.space_after = Pt(4)
        code_style.paragraph_format.left_indent = Cm(0.5)

    # Create Note styles with appropriate colors
    note_configs = {
        'NoteInfo': RGBColor(0x1E, 0x40, 0xAF),      # Blue
        'NoteWarning': RGBColor(0xDC, 0x26, 0x26),   # Red
        'NoteTip': RGBColor(0x05, 0x96, 0x69),       # Green
    }

    for style_name, color in note_configs.items():
        try:
            styles[style_name]
        except KeyError:
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = BODY_FONT
            style.font.size = Pt(10)
            style.font.color.rgb = color
            style.paragraph_format.left_indent = Cm(0.5)
            style.paragraph_format.space_before = Pt(6)
            style.paragraph_format.space_after = Pt(6)
            style.paragraph_format.keep_together = True


def add_jinja(doc: Document, text: str) -> None:
    """Add a Jinja control tag with zero spacing."""
    para = doc.add_paragraph(text, style='JinjaControl')
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)


def set_cell_shading(cell, color_hex: str):
    """Set background shading for a table cell."""
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_document_header_with_metadata(doc: Document) -> None:
    """Add vDocs logo, branding, and metadata to the Word HEADER section.

    Layout (1/5 logo, 4/5 content):
    ┌───────┬────────────────────────────────────────┐
    │       │  vDocs                                 │
    │ LOGO  │  ┌────────────┬──────────────────────┐ │
    │       │  │ Date: ...  │ Language: ...        │ │
    │       │  │ Doc ID: ...                       │ │
    │       │  └────────────┴──────────────────────┘ │
    └───────┴────────────────────────────────────────┘
    """
    from docx.oxml import parse_xml
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT

    section = doc.sections[0]
    header = section.header

    # Clear default paragraph
    header.paragraphs[0].clear()

    # Create main 2-column layout table in HEADER section
    main_table = header.add_table(rows=1, cols=2, width=Cm(16.0))
    main_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    main_table.autofit = False

    # Set column widths - 1/5 for logo, 4/5 for content (total ~16cm usable width)
    main_table.columns[0].width = Cm(3.2)   # Logo column (1/5)
    main_table.columns[1].width = Cm(12.8)  # Content column (4/5)

    # Also set cell widths explicitly to prevent auto-resize
    for row in main_table.rows:
        row.cells[0].width = Cm(3.2)
        row.cells[1].width = Cm(12.8)

    # Remove table borders (invisible layout table)
    tbl = main_table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    tblBorders = parse_xml(
        r'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        r'<w:top w:val="nil"/><w:left w:val="nil"/><w:bottom w:val="nil"/>'
        r'<w:right w:val="nil"/><w:insideH w:val="nil"/><w:insideV w:val="nil"/>'
        r'</w:tblBorders>'
    )
    tblPr.append(tblBorders)

    # ─────────────────────────────────────────────────────────────
    # LEFT COLUMN: Logo (2.2cm)
    # ─────────────────────────────────────────────────────────────
    logo_cell = main_table.cell(0, 0)
    logo_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    logo_para = logo_cell.paragraphs[0]
    logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if LOGO_PATH.exists():
        run = logo_para.add_run()
        run.add_picture(str(LOGO_PATH), width=Cm(2.2), height=Cm(2.2))

    # ─────────────────────────────────────────────────────────────
    # RIGHT COLUMN: vDocs title + metadata table
    # ─────────────────────────────────────────────────────────────
    content_cell = main_table.cell(0, 1)
    content_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    # vDocs title
    title_para = content_cell.paragraphs[0]
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run_v = title_para.add_run("v")
    run_v.font.size = Pt(27)
    run_v.font.bold = True
    run_v.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)  # Dark blue

    run_d = title_para.add_run("D")
    run_d.font.size = Pt(27)
    run_d.font.bold = True
    run_d.font.color.rgb = RGBColor(0x22, 0x8B, 0xE6)  # Primary blue (brand color)

    run_ocs = title_para.add_run("ocs")
    run_ocs.font.size = Pt(27)
    run_ocs.font.bold = True
    run_ocs.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)  # Dark blue

    # Add metadata table inside the right cell
    meta_table = content_cell.add_table(rows=2, cols=2)
    meta_table.autofit = False

    # Set metadata table column widths
    meta_table.columns[0].width = Cm(5.5)
    meta_table.columns[1].width = Cm(5.5)

    # Style metadata table with subtle borders
    meta_tbl = meta_table._tbl
    meta_tblPr = meta_tbl.tblPr if meta_tbl.tblPr is not None else parse_xml(r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    meta_borders = parse_xml(
        r'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        r'<w:top w:val="single" w:sz="4" w:color="E2E8F0"/>'
        r'<w:left w:val="single" w:sz="4" w:color="E2E8F0"/>'
        r'<w:bottom w:val="single" w:sz="4" w:color="E2E8F0"/>'
        r'<w:right w:val="single" w:sz="4" w:color="E2E8F0"/>'
        r'<w:insideH w:val="single" w:sz="4" w:color="E2E8F0"/>'
        r'<w:insideV w:val="single" w:sz="4" w:color="E2E8F0"/>'
        r'</w:tblBorders>'
    )
    meta_tblPr.append(meta_borders)

    # Row 1: Date and Language
    for i, (label, value) in enumerate([
        ("📅 Date", "{{ generated_date }}"),
        ("🌐 Language", "{{ language_upper }}")
    ]):
        cell = meta_table.cell(0, i)
        set_cell_shading(cell, "F8FAFC")
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(f"{label}: ")
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        val_run = para.add_run(value)
        val_run.font.size = Pt(9)
        val_run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    # Row 2: Doc ID (merged across both columns)
    id_cell = meta_table.cell(1, 0)
    id_cell.merge(meta_table.cell(1, 1))
    set_cell_shading(id_cell, "F8FAFC")
    id_para = id_cell.paragraphs[0]
    id_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    id_run = id_para.add_run("📋 Doc ID: ")
    id_run.bold = True
    id_run.font.size = Pt(9)
    id_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    id_val = id_para.add_run("{{ doc_id }}")
    id_val.font.size = Pt(9)
    id_val.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)


def add_header(doc: Document, subtitle: str = "") -> None:
    """Add common header section to template."""
    # Add Word document header with logo and metadata
    add_document_header_with_metadata(doc)

    # Breathing space
    doc.add_paragraph()

    # Title
    title = doc.add_paragraph("{{ title }}", style='Title')
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()


def create_step_manual_template(output_path: Path) -> None:
    """Create template for step-by-step manual format."""
    doc = Document()
    setup_base_styles(doc)
    set_document_format_in_docx(doc, "step-manual")
    add_header(doc, "Step-by-step Manual")

    # Introduction
    add_jinja(doc, "{% if introduction %}")
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("{{ introduction }}")
    add_jinja(doc, "{% endif %}")

    # Steps - each step on a new page (except first)
    add_jinja(doc, "{% for step in semantic_steps %}")
    add_jinja(doc, "{% if not loop.first %}{{ page_break }}{% endif %}")
    doc.add_heading("Step {{ step.number }}: {{ step.title }}", level=2)
    doc.add_paragraph("{{ step.description }}")

    # Step image
    add_jinja(doc, "{% if step.image %}")
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.add_run("{{ step.image }}")
    add_jinja(doc, "{% endif %}")

    add_jinja(doc, "{% endfor %}")

    # Notes
    add_jinja(doc, "{% for note in notes %}")
    add_jinja(doc, "{% if note.type == 'warning' %}")
    doc.add_paragraph("⚠️ Warning: {{ note.content }}", style='NoteWarning')
    add_jinja(doc, "{% elif note.type == 'tip' %}")
    doc.add_paragraph("💡 Tip: {{ note.content }}", style='NoteTip')
    add_jinja(doc, "{% else %}")
    doc.add_paragraph("ℹ️ {{ note.content }}", style='NoteInfo')
    add_jinja(doc, "{% endif %}")
    add_jinja(doc, "{% endfor %}")

    # Conclusion
    add_jinja(doc, "{% if conclusion %}")
    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph("{{ conclusion }}")
    add_jinja(doc, "{% endif %}")

    doc.save(output_path)
    print(f"  ✓ Created: {output_path.name}")


def create_quick_guide_template(output_path: Path) -> None:
    """Create template for quick guide format."""
    doc = Document()
    setup_base_styles(doc)
    set_document_format_in_docx(doc, "quick-guide")
    add_header(doc, "Quick Guide")

    # Overview
    add_jinja(doc, "{% if overview %}")
    doc.add_heading("Overview", level=1)
    doc.add_paragraph("{{ overview }}")
    add_jinja(doc, "{% endif %}")

    # Key Points
    add_jinja(doc, "{% if keypoints %}")
    doc.add_heading("Key Points", level=1)

    add_jinja(doc, "{% for kp in keypoints %}")
    add_jinja(doc, "{% if kp.title %}")
    doc.add_heading("{{ kp.title }}", level=2)
    add_jinja(doc, "{% endif %}")
    doc.add_paragraph("{{ kp.content }}")

    # Keypoint image
    add_jinja(doc, "{% if kp.image %}")
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.add_run("{{ kp.image }}")
    add_jinja(doc, "{% endif %}")

    add_jinja(doc, "{% endfor %}")
    add_jinja(doc, "{% endif %}")

    # Tips
    add_jinja(doc, "{% if tips %}")
    doc.add_heading("Tips", level=1)
    add_jinja(doc, "{% for tip in tips %}")
    doc.add_paragraph("💡 {{ tip.content }}", style='NoteTip')

    # Tip image
    add_jinja(doc, "{% if tip.image %}")
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.add_run("{{ tip.image }}")
    add_jinja(doc, "{% endif %}")

    add_jinja(doc, "{% endfor %}")
    add_jinja(doc, "{% endif %}")

    doc.save(output_path)
    print(f"  ✓ Created: {output_path.name}")


def create_reference_template(output_path: Path) -> None:
    """Create template for reference document format."""
    doc = Document()
    setup_base_styles(doc)
    set_document_format_in_docx(doc, "reference")
    add_header(doc, "Reference Document")

    # Sections
    add_jinja(doc, "{% for section in sections %}")
    add_jinja(doc, "{% if section.title %}")
    doc.add_heading("{{ section.title }}", level=1)
    add_jinja(doc, "{% endif %}")
    doc.add_paragraph("{{ section.content }}")

    # Section image
    add_jinja(doc, "{% if section.image %}")
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.add_run("{{ section.image }}")
    add_jinja(doc, "{% endif %}")

    add_jinja(doc, "{% endfor %}")

    # Definitions
    add_jinja(doc, "{% if definitions %}")
    doc.add_heading("Definitions", level=1)
    add_jinja(doc, "{% for def in definitions %}")
    term_para = doc.add_paragraph()
    term_run = term_para.add_run("{{ def.term }}: ")
    term_run.bold = True
    term_para.add_run("{{ def.content }}")
    add_jinja(doc, "{% endfor %}")
    add_jinja(doc, "{% endif %}")

    # Examples
    add_jinja(doc, "{% if examples %}")
    doc.add_heading("Examples", level=1)
    add_jinja(doc, "{% for ex in examples %}")
    add_jinja(doc, "{% if ex.title %}")
    doc.add_heading("{{ ex.title }}", level=2)
    add_jinja(doc, "{% endif %}")
    doc.add_paragraph("{{ ex.content }}")

    # Example image
    add_jinja(doc, "{% if ex.image %}")
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.add_run("{{ ex.image }}")
    add_jinja(doc, "{% endif %}")

    add_jinja(doc, "{% endfor %}")
    add_jinja(doc, "{% endif %}")

    doc.save(output_path)
    print(f"  ✓ Created: {output_path.name}")


def create_summary_template(output_path: Path) -> None:
    """Create template for executive summary format."""
    doc = Document()
    setup_base_styles(doc)
    set_document_format_in_docx(doc, "summary")
    add_header(doc, "Executive Summary")

    # Highlights
    add_jinja(doc, "{% if highlights %}")
    doc.add_heading("Key Highlights", level=1)
    doc.add_paragraph("{{ highlights }}")
    add_jinja(doc, "{% endif %}")

    # Findings
    add_jinja(doc, "{% if findings %}")
    doc.add_heading("Findings", level=1)
    add_jinja(doc, "{% for finding in findings %}")
    add_jinja(doc, "{% if finding.title %}")
    doc.add_heading("{{ finding.title }}", level=2)
    add_jinja(doc, "{% else %}")
    doc.add_heading("Finding {{ finding.number }}", level=2)
    add_jinja(doc, "{% endif %}")
    doc.add_paragraph("{{ finding.content }}")

    # Finding image
    add_jinja(doc, "{% if finding.image %}")
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.add_run("{{ finding.image }}")
    add_jinja(doc, "{% endif %}")

    add_jinja(doc, "{% endfor %}")
    add_jinja(doc, "{% endif %}")

    # Recommendations
    add_jinja(doc, "{% if recommendations %}")
    doc.add_heading("Recommendations", level=1)
    add_jinja(doc, "{% for rec in recommendations %}")
    add_jinja(doc, "{% if rec.title %}")
    doc.add_heading("{{ rec.title }}", level=2)
    add_jinja(doc, "{% else %}")
    doc.add_heading("Recommendation {{ rec.number }}", level=2)
    add_jinja(doc, "{% endif %}")
    doc.add_paragraph("{{ rec.content }}")
    add_jinja(doc, "{% endfor %}")
    add_jinja(doc, "{% endif %}")

    doc.save(output_path)
    print(f"  ✓ Created: {output_path.name}")


def create_legacy_template(output_path: Path) -> None:
    """Create a legacy template for manuals without semantic tags."""
    doc = Document()
    setup_base_styles(doc)
    add_header(doc, "Manual")

    # Steps loop (legacy format using markdown-parsed steps)
    add_jinja(doc, "{% for step in steps %}")
    doc.add_heading("Step {{ step.number }}{% if step.title %}: {{ step.title }}{% endif %}", level=2)
    doc.add_paragraph("{{ step.description }}")

    add_jinja(doc, "{% if step.has_image %}")
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.add_run("{{ step.image }}")
    add_jinja(doc, "{% endif %}")

    add_jinja(doc, "{% endfor %}")

    doc.save(output_path)
    print(f"  ✓ Created: {output_path.name}")


def set_document_format_in_docx(doc: Document, document_format: str) -> None:
    """Store document format in the docx core properties (category field)."""
    doc.core_properties.category = f"vdocs:{document_format}"


def create_template_metadata(template_path: Path, document_format: str) -> None:
    """Create metadata JSON file for a template."""
    import json
    meta_path = template_path.with_suffix(".json")
    metadata = {
        "document_format": document_format,
        "is_default": True,
    }
    with open(meta_path, "w", encoding="utf-8") as f:
        json.dump(metadata, f, indent=2)


def main():
    """Create all default templates."""
    script_dir = Path(__file__).parent
    project_root = script_dir.parent

    # Output to root /templates/ (committed to git)
    templates_dir = project_root / "templates"
    templates_dir.mkdir(parents=True, exist_ok=True)

    print(f"Creating default templates in: {templates_dir}\n")

    # Map template names to their creation functions and document formats
    templates = {
        "step-manual": (create_step_manual_template, "step-manual"),
        "quick-guide": (create_quick_guide_template, "quick-guide"),
        "reference": (create_reference_template, "reference"),
        "summary": (create_summary_template, "summary"),
        "default-manual": (create_legacy_template, None),  # Legacy has no specific format
    }

    for name, (create_func, doc_format) in templates.items():
        template_path = templates_dir / f"{name}.docx"
        create_func(template_path)
        if doc_format:
            create_template_metadata(template_path, doc_format)

    print("\n✅ All templates created successfully!")


if __name__ == "__main__":
    main()
