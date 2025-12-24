"""Create Word templates for report document formats.

This script generates DOCX templates with Jinja2 placeholders for:
- Incident Report
- Inspection Report
- Progress Report
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Logo path for header
LOGO_PATH = Path(__file__).parent.parent / "data" / "templates" / "vdocs-logo.png"


def add_document_header_with_metadata(doc: Document) -> None:
    """Add vDocs logo, branding, and metadata to the Word HEADER section.

    Layout (1/5 logo, 4/5 content):
    ┌───────┬────────────────────────────────────────┐
    │       │  vDocs                                 │
    │ LOGO  │  ┌────────────┬──────────────────────┐ │
    │       │  │ Date: ...  │ Language: ...        │ │
    │       │  │ Report ID: ...                    │ │
    │       │  └────────────┴──────────────────────┘ │
    └───────┴────────────────────────────────────────┘
    """
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    from docx.shared import Cm
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT

    section = doc.sections[0]
    header = section.header

    # Clear default paragraph
    header.paragraphs[0].clear()

    # Create main 2-column layout table in HEADER section
    main_table = header.add_table(rows=1, cols=2, width=Cm(16.0))
    main_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    main_table.autofit = False

    # Set column widths - 1/5 for logo, 4/5 for content (total ~16cm usable width)
    main_table.columns[0].width = Cm(3.2)   # Logo column (1/5)
    main_table.columns[1].width = Cm(12.8)  # Content column (4/5)

    # Also set cell widths explicitly to prevent auto-resize
    for row in main_table.rows:
        row.cells[0].width = Cm(3.2)
        row.cells[1].width = Cm(12.8)

    # Remove table borders (invisible layout table)
    tbl = main_table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    tblBorders = parse_xml(
        r'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        r'<w:top w:val="nil"/><w:left w:val="nil"/><w:bottom w:val="nil"/>'
        r'<w:right w:val="nil"/><w:insideH w:val="nil"/><w:insideV w:val="nil"/>'
        r'</w:tblBorders>'
    )
    tblPr.append(tblBorders)

    # ─────────────────────────────────────────────────────────────
    # LEFT COLUMN: Logo (2.2cm)
    # ─────────────────────────────────────────────────────────────
    logo_cell = main_table.cell(0, 0)
    logo_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    logo_para = logo_cell.paragraphs[0]
    logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if LOGO_PATH.exists():
        run = logo_para.add_run()
        run.add_picture(str(LOGO_PATH), width=Cm(2.2), height=Cm(2.2))

    # ─────────────────────────────────────────────────────────────
    # RIGHT COLUMN: vDocs title + metadata table
    # ─────────────────────────────────────────────────────────────
    content_cell = main_table.cell(0, 1)
    content_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    # vDocs title
    title_para = content_cell.paragraphs[0]
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run_v = title_para.add_run("v")
    run_v.font.size = Pt(27)
    run_v.font.bold = True
    run_v.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)  # Dark blue

    run_d = title_para.add_run("D")
    run_d.font.size = Pt(27)
    run_d.font.bold = True
    run_d.font.color.rgb = RGBColor(0x22, 0x8B, 0xE6)  # Primary blue (brand color)

    run_ocs = title_para.add_run("ocs")
    run_ocs.font.size = Pt(27)
    run_ocs.font.bold = True
    run_ocs.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)  # Dark blue

    # Add metadata table inside the right cell
    meta_table = content_cell.add_table(rows=2, cols=2)
    meta_table.autofit = False

    # Set metadata table column widths
    meta_table.columns[0].width = Cm(5.5)
    meta_table.columns[1].width = Cm(5.5)

    # Style metadata table with subtle borders
    meta_tbl = meta_table._tbl
    meta_tblPr = meta_tbl.tblPr if meta_tbl.tblPr is not None else parse_xml(r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    meta_borders = parse_xml(
        r'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        r'<w:top w:val="single" w:sz="4" w:color="E2E8F0"/>'
        r'<w:left w:val="single" w:sz="4" w:color="E2E8F0"/>'
        r'<w:bottom w:val="single" w:sz="4" w:color="E2E8F0"/>'
        r'<w:right w:val="single" w:sz="4" w:color="E2E8F0"/>'
        r'<w:insideH w:val="single" w:sz="4" w:color="E2E8F0"/>'
        r'<w:insideV w:val="single" w:sz="4" w:color="E2E8F0"/>'
        r'</w:tblBorders>'
    )
    meta_tblPr.append(meta_borders)

    # Row 1: Date and Language
    for i, (label, value) in enumerate([
        ("📅 Date", "{{ generated_date }}"),
        ("🌐 Language", "{{ language_upper }}")
    ]):
        cell = meta_table.cell(0, i)
        set_cell_shading(cell, "F8FAFC")
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(f"{label}: ")
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        val_run = para.add_run(value)
        val_run.font.size = Pt(9)
        val_run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    # Row 2: Report ID (merged across both columns)
    id_cell = meta_table.cell(1, 0)
    id_cell.merge(meta_table.cell(1, 1))
    set_cell_shading(id_cell, "F8FAFC")
    id_para = id_cell.paragraphs[0]
    id_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    id_run = id_para.add_run("📋 Report ID: ")
    id_run.bold = True
    id_run.font.size = Pt(9)
    id_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    id_val = id_para.add_run("{{ doc_id }}")
    id_val.font.size = Pt(9)
    id_val.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)


def add_document_header(doc: Document) -> None:
    """Legacy: Add vDocs logo and branding to the Word document header section."""
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    from docx.shared import Cm
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

    section = doc.sections[0]
    header = section.header

    # Clear default paragraph
    header.paragraphs[0].clear()

    # Create a table for layout: [Logo] | [vDocs text] - aligned at bottom
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    table.autofit = False

    # Set column widths
    table.columns[0].width = Cm(2.0)   # Logo column
    table.columns[1].width = Cm(3.5)   # Text column

    # Remove table borders (invisible layout table)
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    tblBorders = parse_xml(
        r'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        r'<w:top w:val="nil"/><w:left w:val="nil"/><w:bottom w:val="nil"/>'
        r'<w:right w:val="nil"/><w:insideH w:val="nil"/><w:insideV w:val="nil"/>'
        r'</w:tblBorders>'
    )
    tblPr.append(tblBorders)

    # Logo cell - aligned bottom
    logo_cell = table.cell(0, 0)
    logo_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
    logo_para = logo_cell.paragraphs[0]
    logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add logo (1.73cm x 1.73cm)
    if LOGO_PATH.exists():
        run = logo_para.add_run()
        run.add_picture(str(LOGO_PATH), width=Cm(1.73), height=Cm(1.73))

    # Text cell - aligned bottom
    text_cell = table.cell(0, 1)
    text_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
    text_para = text_cell.paragraphs[0]
    text_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add "vDocs" text with styled "D" - 27pt
    run_v = text_para.add_run("v")
    run_v.font.size = Pt(27)
    run_v.font.bold = True
    run_v.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)  # Dark blue

    run_d = text_para.add_run("D")
    run_d.font.size = Pt(27)
    run_d.font.bold = True
    run_d.font.color.rgb = RGBColor(0x22, 0x8B, 0xE6)  # Primary blue (brand color)

    run_ocs = text_para.add_run("ocs")
    run_ocs.font.size = Pt(27)
    run_ocs.font.bold = True
    run_ocs.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)  # Dark blue


def create_base_styles(doc: Document):
    """Add common styles to the document with modern typography."""
    from docx.shared import Cm

    styles = doc.styles

    # Modern font choices
    HEADING_FONT = "Segoe UI"  # Clean, modern sans-serif
    BODY_FONT = "Segoe UI"     # Consistent look
    # Fallback fonts if Segoe UI not available: Arial, Calibri

    # Normal/body text style
    normal_style = styles['Normal']
    normal_style.font.name = BODY_FONT
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.paragraph_format.line_spacing = 1.15

    # Title style - Bold and impactful
    title_style = styles['Title']
    title_style.font.name = HEADING_FONT
    title_style.font.size = Pt(28)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    title_style.paragraph_format.space_after = Pt(12)

    # Heading 1 - Section headers
    h1_style = styles['Heading 1']
    h1_style.font.name = HEADING_FONT
    h1_style.font.size = Pt(18)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)  # Primary blue
    h1_style.paragraph_format.space_before = Pt(18)
    h1_style.paragraph_format.space_after = Pt(8)

    # Heading 2 - Subsection headers
    h2_style = styles['Heading 2']
    h2_style.font.name = HEADING_FONT
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
    h2_style.paragraph_format.space_before = Pt(12)
    h2_style.paragraph_format.space_after = Pt(6)

    # Heading 3 - Minor headers
    h3_style = styles['Heading 3']
    h3_style.font.name = HEADING_FONT
    h3_style.font.size = Pt(12)
    h3_style.font.bold = True
    h3_style.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
    h3_style.paragraph_format.space_before = Pt(8)
    h3_style.paragraph_format.space_after = Pt(4)


def set_cell_shading(cell, color_hex: str):
    """Set background shading for a table cell."""
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_table_borders(table, color: str = "CCCCCC", size: int = 4):
    """Set borders on a table."""
    from docx.oxml import parse_xml
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        tbl.insert(0, tblPr)

    borders = parse_xml(
        f'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:top w:val="single" w:sz="{size}" w:color="{color}"/>'
        f'<w:left w:val="single" w:sz="{size}" w:color="{color}"/>'
        f'<w:bottom w:val="single" w:sz="{size}" w:color="{color}"/>'
        f'<w:right w:val="single" w:sz="{size}" w:color="{color}"/>'
        f'<w:insideH w:val="single" w:sz="{size}" w:color="{color}"/>'
        f'<w:insideV w:val="single" w:sz="{size}" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    # Remove existing borders if any
    for existing in tblPr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders'):
        tblPr.remove(existing)
    tblPr.append(borders)


def remove_table_borders(table):
    """Remove all borders from a table."""
    from docx.oxml import parse_xml
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        tbl.insert(0, tblPr)

    borders = parse_xml(
        r'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        r'<w:top w:val="nil"/><w:left w:val="nil"/><w:bottom w:val="nil"/>'
        r'<w:right w:val="nil"/><w:insideH w:val="nil"/><w:insideV w:val="nil"/>'
        r'</w:tblBorders>'
    )
    for existing in tblPr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders'):
        tblPr.remove(existing)
    tblPr.append(borders)


def keep_table_together(table):
    """Prevent a table from splitting across pages."""
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.keep_together = True
                para.paragraph_format.keep_with_next = True
    # Last paragraph shouldn't keep_with_next
    last_cell = table.rows[-1].cells[-1]
    if last_cell.paragraphs:
        last_cell.paragraphs[-1].paragraph_format.keep_with_next = False


def add_page_borders(doc: Document, color: str = "1E40AF", size: int = 12, space: int = 24):
    """Add decorative page borders to the document.

    Args:
        doc: The document to add borders to
        color: Border color in hex (without #)
        size: Border width in eighths of a point (12 = 1.5pt)
        space: Space between border and page content in points
    """
    from docx.oxml import parse_xml
    from docx.oxml.ns import qn

    for section in doc.sections:
        sectPr = section._sectPr

        # Create page borders element
        pgBorders = parse_xml(
            f'<w:pgBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:offsetFrom="page">'
            f'<w:top w:val="single" w:sz="{size}" w:space="{space}" w:color="{color}"/>'
            f'<w:left w:val="single" w:sz="{size}" w:space="{space}" w:color="{color}"/>'
            f'<w:bottom w:val="single" w:sz="{size}" w:space="{space}" w:color="{color}"/>'
            f'<w:right w:val="single" w:sz="{size}" w:space="{space}" w:color="{color}"/>'
            f'</w:pgBorders>'
        )

        # Remove existing page borders if any
        for existing in sectPr.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pgBorders'):
            sectPr.remove(existing)

        # Insert after pgSz (page size) or at beginning
        pgSz = sectPr.find(qn('w:pgSz'))
        if pgSz is not None:
            pgSz.addnext(pgBorders)
        else:
            sectPr.insert(0, pgBorders)


def create_incident_report_template(output_path: Path):
    """Create template for Incident Report format with table-based layouts."""
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import parse_xml

    doc = Document()
    create_base_styles(doc)

    # ═══════════════════════════════════════════════════════════════
    # HEADER BLOCK - Logo + vDocs + Metadata in 2-column layout
    # ═══════════════════════════════════════════════════════════════
    add_document_header_with_metadata(doc)

    # ═══════════════════════════════════════════════════════════════
    # BREATHING SPACE - Line break between header and title
    # ═══════════════════════════════════════════════════════════════
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════
    # TITLE SECTION - Large title with underline accent
    # ═══════════════════════════════════════════════════════════════
    title = doc.add_heading('{{ title }}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════
    # SUMMARY SECTION - Highlighted box
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('Executive Summary', level=1)

    summary_table = doc.add_table(rows=1, cols=1)
    set_table_borders(summary_table, color="3B82F6", size=8)
    summary_cell = summary_table.cell(0, 0)
    set_cell_shading(summary_cell, "EFF6FF")

    summary_para = summary_cell.paragraphs[0]
    summary_para.add_run('{{ summary }}')

    # Summary image if exists
    doc.add_paragraph('{% if summary_image %}')
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.add_run('{{ summary_image }}')
    doc.add_paragraph('{% endif %}')

    # ═══════════════════════════════════════════════════════════════
    # LOCATION SECTION - Two-row layout (text, then image)
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('Location Details', level=1)

    doc.add_paragraph('{{ location }}')

    # Image on its own row
    doc.add_paragraph('{% if location_image %}')
    loc_img_para = doc.add_paragraph()
    loc_img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    loc_img_para.add_run('{{ location_image }}')
    doc.add_paragraph('{% endif %}')

    # ═══════════════════════════════════════════════════════════════
    # DETAILED FINDINGS - Card-style layout (title, content, image rows)
    # Each finding starts on a new page
    # ═══════════════════════════════════════════════════════════════
    findings_heading = doc.add_heading('Detailed Findings', level=1)
    # Keep heading with first content
    findings_heading.paragraph_format.keep_with_next = True

    # Jinja loop for findings - page break before each (except first)
    doc.add_paragraph('{% for finding in findings %}')
    doc.add_paragraph('{% if not loop.first %}{{ page_break }}{% endif %}')

    # Finding card table - 3 rows: title, content, image
    finding_table = doc.add_table(rows=3, cols=1)
    set_table_borders(finding_table, color="E2E8F0", size=4)
    finding_table.columns[0].width = Inches(6.5)
    keep_table_together(finding_table)

    # Row 1: Title (with top spacing for breathing room after page breaks)
    title_cell = finding_table.cell(0, 0)
    set_cell_shading(title_cell, "F1F5F9")
    title_para = title_cell.paragraphs[0]
    title_para.paragraph_format.space_before = Pt(12)  # Top spacing within cell
    title_run = title_para.add_run('🔍 {{ finding.title }}')
    title_run.bold = True
    title_run.font.size = Pt(11)
    title_run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    # Row 2: Content
    content_cell = finding_table.cell(1, 0)
    content_cell.paragraphs[0].add_run('{{ finding.content }}')

    # Row 3: Image (centered)
    img_cell = finding_table.cell(2, 0)
    img_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_cell.paragraphs[0].add_run('{% if finding.has_image %}{{ finding.image }}{% endif %}')

    doc.add_paragraph()  # Spacing between cards
    doc.add_paragraph('{% endfor %}')

    # ═══════════════════════════════════════════════════════════════
    # EVIDENCE SECTION - Card-style (title, content, image rows)
    # Each evidence item starts on a new page
    # ═══════════════════════════════════════════════════════════════
    # Page break before Evidence section to separate from Findings
    doc.add_paragraph('{{ page_break }}')

    evidence_heading = doc.add_heading('Evidence', level=1)
    # Keep heading with first content
    evidence_heading.paragraph_format.keep_with_next = True

    # Jinja loop for evidence - page break before each (except first)
    doc.add_paragraph('{% for evidence in evidences %}')
    doc.add_paragraph('{% if not loop.first %}{{ page_break }}{% endif %}')

    # Evidence card table - 3 rows: title, content, image
    evidence_table = doc.add_table(rows=3, cols=1)
    set_table_borders(evidence_table, color="D1D5DB", size=4)
    evidence_table.columns[0].width = Inches(6.5)
    keep_table_together(evidence_table)

    # Row 1: Title (with extra top spacing for breathing room from header)
    ev_title_cell = evidence_table.cell(0, 0)
    set_cell_shading(ev_title_cell, "F9FAFB")
    ev_title_para = ev_title_cell.paragraphs[0]
    ev_title_para.paragraph_format.space_before = Pt(12)  # Top spacing within cell
    ev_title_run = ev_title_para.add_run('📷 {{ evidence.title }}')
    ev_title_run.bold = True
    ev_title_run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

    # Row 2: Content
    ev_content_cell = evidence_table.cell(1, 0)
    ev_content_cell.paragraphs[0].add_run('{{ evidence.content }}')

    # Row 3: Image (centered)
    ev_img_cell = evidence_table.cell(2, 0)
    ev_img_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    ev_img_cell.paragraphs[0].add_run('{% if evidence.has_image %}{{ evidence.image }}{% endif %}')

    doc.add_paragraph()
    doc.add_paragraph('{% endfor %}')

    # ═══════════════════════════════════════════════════════════════
    # SEVERITY ASSESSMENT - Visual indicator box (on its own page)
    # ═══════════════════════════════════════════════════════════════
    doc.add_paragraph('{{ page_break }}')
    doc.add_heading('Severity Assessment', level=1)

    severity_table = doc.add_table(rows=1, cols=2)
    set_table_borders(severity_table, color="DC2626", size=6)
    severity_table.columns[0].width = Inches(1.5)
    severity_table.columns[1].width = Inches(5.0)

    # Severity badge
    sev_badge_cell = severity_table.cell(0, 0)
    set_cell_shading(sev_badge_cell, "FEF2F2")
    sev_badge_para = sev_badge_cell.paragraphs[0]
    sev_badge_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sev_run = sev_badge_para.add_run('⚠️ {{ severity.level|upper }}')
    sev_run.bold = True
    sev_run.font.size = Pt(14)
    sev_run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

    # Severity description
    sev_desc_cell = severity_table.cell(0, 1)
    sev_desc_cell.paragraphs[0].add_run('{{ severity.content }}')

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════
    # RECOMMENDATIONS - Priority-coded cards (on its own page)
    # ═══════════════════════════════════════════════════════════════
    doc.add_paragraph('{{ page_break }}')
    doc.add_heading('Recommendations', level=1)

    doc.add_paragraph('{% for rec in recommendations %}')

    rec_table = doc.add_table(rows=1, cols=1)
    set_table_borders(rec_table, color="10B981", size=4)
    rec_cell = rec_table.cell(0, 0)

    # Title with priority badge
    rec_title_para = rec_cell.paragraphs[0]
    rec_title_run = rec_title_para.add_run('✅ {{ rec.title }}')
    rec_title_run.bold = True
    rec_title_run.font.color.rgb = RGBColor(0x05, 0x96, 0x69)

    # Priority line
    priority_para = rec_cell.add_paragraph()
    priority_run = priority_para.add_run('Priority: {{ rec.priority|title }}')
    priority_run.italic = True
    priority_run.font.size = Pt(9)
    priority_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    # Content
    rec_cell.add_paragraph('{{ rec.content }}')

    doc.add_paragraph()
    doc.add_paragraph('{% endfor %}')

    # ═══════════════════════════════════════════════════════════════
    # NEXT STEPS - Checklist style (on its own page)
    # ═══════════════════════════════════════════════════════════════
    doc.add_paragraph('{{ page_break }}')
    doc.add_heading('Next Steps', level=1)

    next_table = doc.add_table(rows=1, cols=1)
    set_table_borders(next_table, color="8B5CF6", size=4)
    next_cell = next_table.cell(0, 0)
    set_cell_shading(next_cell, "F5F3FF")
    next_cell.paragraphs[0].add_run('{{ next_steps }}')

    # ═══════════════════════════════════════════════════════════════
    # FOOTER
    # ═══════════════════════════════════════════════════════════════
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_run = footer_para.add_run('Generated by vDocs • {{ generated_at }}')
    footer_run.italic = True
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Add decorative page borders
    add_page_borders(doc, color="1E40AF", size=12, space=24)

    doc.save(output_path)
    print(f"  ✓ Created: {output_path.name}")


def create_inspection_report_template(output_path: Path):
    """Create template for Inspection Report format."""
    doc = Document()
    create_base_styles(doc)

    # Header with logo, vDocs branding, and metadata
    add_document_header_with_metadata(doc)

    # Breathing space
    doc.add_paragraph()

    # Title
    title = doc.add_heading('{{ title }}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()

    # Overview section
    doc.add_heading('Inspection Overview', level=1)
    doc.add_paragraph('{{ overview }}')

    # Inspection Items
    doc.add_heading('Inspection Items', level=1)
    doc.add_paragraph('''{% for item in inspection_items %}
{{ item.title }}
Status: {{ item.status|upper }}
{{ item.content }}
{% if item.has_image %}{{ item.image }}{% endif %}
---
{% endfor %}''')

    # Findings
    doc.add_heading('Findings', level=1)
    doc.add_paragraph('''{% for finding in findings %}
{{ finding.title }}
{{ finding.content }}
{% if finding.has_image %}{{ finding.image }}{% endif %}
{% endfor %}''')

    # Status Summary
    doc.add_heading('Overall Status Summary', level=1)
    doc.add_paragraph('{{ status_summary }}')

    # Recommendations
    doc.add_heading('Recommendations', level=1)
    doc.add_paragraph('''{% for rec in recommendations %}
{{ rec.title }}
Priority: {{ rec.priority|title }}
{{ rec.content }}
{% endfor %}''')

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Generated: {{ generated_at }}').italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.save(output_path)
    print(f"Created: {output_path}")


def create_progress_report_template(output_path: Path):
    """Create template for Progress Report format."""
    doc = Document()
    create_base_styles(doc)

    # Header with logo, vDocs branding, and metadata
    add_document_header_with_metadata(doc)

    # Breathing space
    doc.add_paragraph()

    # Title
    title = doc.add_heading('{{ title }}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()

    # Reporting Period
    doc.add_heading('Reporting Period', level=1)
    doc.add_paragraph('{{ period }}')
    doc.add_paragraph('{% if period_image %}{{ period_image }}{% endif %}')

    # Accomplishments
    doc.add_heading('Accomplishments', level=1)
    doc.add_paragraph('''{% for acc in accomplishments %}
{{ acc.title }}
{{ acc.content }}
{% if acc.has_image %}{{ acc.image }}{% endif %}
---
{% endfor %}''')

    # Issues
    doc.add_heading('Issues & Challenges', level=1)
    doc.add_paragraph('''{% for issue in issues %}
{{ issue.title }}
Impact: {{ issue.impact|title }}
{{ issue.content }}
{% if issue.has_image %}{{ issue.image }}{% endif %}
---
{% endfor %}''')

    # Next Steps
    doc.add_heading('Next Steps', level=1)
    doc.add_paragraph('{{ next_steps }}')

    # Timeline
    doc.add_heading('Timeline Status', level=1)
    doc.add_paragraph('{{ timeline }}')

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Generated: {{ generated_at }}').italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.save(output_path)
    print(f"Created: {output_path}")


def main():
    """Generate all report templates."""
    templates_dir = Path(__file__).parent.parent / "data" / "templates"
    templates_dir.mkdir(parents=True, exist_ok=True)

    create_incident_report_template(templates_dir / "incident-report.docx")
    create_inspection_report_template(templates_dir / "inspection-report.docx")
    create_progress_report_template(templates_dir / "progress-report.docx")

    print("\nAll report templates created successfully!")


if __name__ == "__main__":
    main()
